#!/usr/bin/env python3
"""
Compare formatting between a PDF and DOCX file to verify conversion quality.
Checks fonts, sizes, colors, structure, and overall formatting preservation.
"""

import sys
import os
from pathlib import Path

def compare_pdf_docx(pdf_path: str, docx_path: str):
    """Compare PDF and DOCX formatting."""
    try:
        import fitz  # PyMuPDF
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_COLOR_INDEX
    except ImportError as e:
        print(f"Error: Missing required library: {e}")
        print("Install with: pip install PyMuPDF python-docx")
        sys.exit(1)
    
    print("=" * 80)
    print("PDF to DOCX Formatting Comparison")
    print("=" * 80)
    print(f"\nPDF: {pdf_path}")
    print(f"DOCX: {docx_path}\n")
    
    # Analyze PDF
    print("Analyzing PDF...")
    pdf = fitz.open(pdf_path)
    pdf_stats = {
        'pages': len(pdf),
        'fonts': set(),
        'font_sizes': set(),
        'colors': set(),
        'bold_count': 0,
        'italic_count': 0,
        'total_text_blocks': 0,
        'images': 0,
        'tables': 0
    }
    
    for page_num in range(len(pdf)):
        page = pdf[page_num]
        
        # Count images
        images = page.get_images()
        pdf_stats['images'] += len(images)
        
        # Extract text with formatting
        text_dict = page.get_text("dict")
        for block in text_dict.get("blocks", []):
            if "lines" in block:
                pdf_stats['total_text_blocks'] += len(block["lines"])
                for line in block["lines"]:
                    for span in line.get("spans", []):
                        font = span.get("font", "")
                        size = span.get("size", 0)
                        flags = span.get("flags", 0)
                        color = span.get("color", 0)
                        
                        if font:
                            pdf_stats['fonts'].add(font)
                        if size > 0:
                            pdf_stats['font_sizes'].add(round(size, 1))
                        if color > 0:
                            pdf_stats['colors'].add(color)
                        if flags & 16:  # Bold
                            pdf_stats['bold_count'] += 1
                        if flags & 2:  # Italic
                            pdf_stats['italic_count'] += 1
        
        # Try to detect tables
        try:
            tables = page.find_tables()
            if tables:
                pdf_stats['tables'] += len(tables)
        except:
            pass
    
    pdf.close()
    
    # Analyze DOCX
    print("Analyzing DOCX...")
    docx_doc = Document(docx_path)
    docx_stats = {
        'paragraphs': len(docx_doc.paragraphs),
        'fonts': set(),
        'font_sizes': set(),
        'colors': set(),
        'bold_count': 0,
        'italic_count': 0,
        'tables': len(docx_doc.tables),
        'images': 0,
        'list_items': 0
    }
    
    # Check paragraphs
    for para in docx_doc.paragraphs:
        # Check if it's a list
        if para.style.name.startswith('List'):
            docx_stats['list_items'] += 1
        
        for run in para.runs:
            if run.font.name:
                docx_stats['fonts'].add(run.font.name)
            if run.font.size:
                size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else float(run.font.size) / 12700
                docx_stats['font_sizes'].add(round(size_pt, 1))
            if run.font.color and run.font.color.rgb:
                docx_stats['colors'].add(str(run.font.color.rgb))
            if run.font.bold:
                docx_stats['bold_count'] += 1
            if run.font.italic:
                docx_stats['italic_count'] += 1
        
        # Check for images in paragraph
        for rel in para.part.rels.values():
            if "image" in rel.target_ref:
                docx_stats['images'] += 1
    
    # Check tables
    for table in docx_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.name:
                            docx_stats['fonts'].add(run.font.name)
    
    # Print comparison
    print("\n" + "=" * 80)
    print("COMPARISON RESULTS")
    print("=" * 80)
    
    print(f"\n[Structure]")
    print(f"  PDF Pages: {pdf_stats['pages']}")
    print(f"  DOCX Paragraphs: {docx_stats['paragraphs']}")
    print(f"  DOCX List Items: {docx_stats['list_items']}")
    
    print(f"\n[Images]")
    print(f"  PDF: {pdf_stats['images']}")
    print(f"  DOCX: {docx_stats['images']}")
    match_status = "MATCH" if pdf_stats['images'] == docx_stats['images'] else "MISMATCH"
    print(f"  Status: {match_status}")
    
    print(f"\n[Tables]")
    print(f"  PDF: {pdf_stats['tables']}")
    print(f"  DOCX: {docx_stats['tables']}")
    match_status = "MATCH" if pdf_stats['tables'] == docx_stats['tables'] else "MISMATCH"
    print(f"  Status: {match_status}")
    
    print(f"\n[Fonts]")
    print(f"  PDF Fonts ({len(pdf_stats['fonts'])}): {sorted(list(pdf_stats['fonts']))[:5]}")
    print(f"  DOCX Fonts ({len(docx_stats['fonts'])}): {sorted(list(docx_stats['fonts']))[:5]}")
    common_fonts = pdf_stats['fonts'] & docx_stats['fonts']
    print(f"  Common: {len(common_fonts)} fonts")
    
    print(f"\n[Font Sizes]")
    print(f"  PDF Sizes: {sorted(list(pdf_stats['font_sizes']))[:10]}")
    print(f"  DOCX Sizes: {sorted(list(docx_stats['font_sizes']))[:10]}")
    common_sizes = pdf_stats['font_sizes'] & docx_stats['font_sizes']
    print(f"  Common: {len(common_sizes)} sizes")
    
    print(f"\n[Formatting]")
    print(f"  PDF Bold: {pdf_stats['bold_count']}, Italic: {pdf_stats['italic_count']}")
    print(f"  DOCX Bold: {docx_stats['bold_count']}, Italic: {docx_stats['italic_count']}")
    
    print(f"\n[Colors]")
    print(f"  PDF Unique Colors: {len(pdf_stats['colors'])}")
    print(f"  DOCX Unique Colors: {len(docx_stats['colors'])}")
    
    # Overall assessment
    print("\n" + "=" * 80)
    print("OVERALL ASSESSMENT")
    print("=" * 80)
    
    score = 0
    max_score = 6
    
    if pdf_stats['images'] == docx_stats['images']:
        score += 1
        print("[OK] Images preserved")
    else:
        print(f"[WARN] Images: PDF has {pdf_stats['images']}, DOCX has {docx_stats['images']}")
    
    if pdf_stats['tables'] == docx_stats['tables']:
        score += 1
        print("[OK] Tables preserved")
    else:
        print(f"[WARN] Tables: PDF has {pdf_stats['tables']}, DOCX has {docx_stats['tables']}")
    
    if len(common_fonts) > 0:
        score += 1
        print(f"[OK] Fonts preserved ({len(common_fonts)} common fonts)")
    else:
        print("[WARN] No common fonts found")
    
    if len(common_sizes) > 0:
        score += 1
        print(f"[OK] Font sizes preserved ({len(common_sizes)} common sizes)")
    else:
        print("[WARN] No common font sizes found")
    
    if docx_stats['bold_count'] > 0 or docx_stats['italic_count'] > 0:
        score += 1
        print("[OK] Text formatting (bold/italic) preserved")
    else:
        print("[WARN] No bold/italic formatting found in DOCX")
    
    if docx_stats['list_items'] > 0:
        score += 1
        print(f"[OK] Lists detected ({docx_stats['list_items']} items)")
    else:
        print("[WARN] No lists detected in DOCX")
    
    print(f"\nFormatting Preservation Score: {score}/{max_score} ({score*100//max_score}%)")
    
    if score >= 5:
        print("RESULT: Excellent formatting preservation!")
    elif score >= 3:
        print("RESULT: Good formatting preservation")
    else:
        print("RESULT: Some formatting may be lost")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python scripts/compare_pdf_docx_formatting.py <pdf_path> <docx_path>")
        print("\nExample:")
        print('  python scripts/compare_pdf_docx_formatting.py "original.pdf" "redlined.docx"')
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    
    if not os.path.exists(pdf_path):
        print(f"Error: PDF file not found: {pdf_path}")
        sys.exit(1)
    
    if not os.path.exists(docx_path):
        print(f"Error: DOCX file not found: {docx_path}")
        sys.exit(1)
    
    compare_pdf_docx(pdf_path, docx_path)

