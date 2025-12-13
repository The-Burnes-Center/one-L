"""
Model integration for Claude 4 Sonnet thinking with AWS Bedrock.
Handles tool calling for comprehensive document review.

Features:
- Detailed logging of LLM thinking process during all inference calls
- Thinking logs are captured and logged for:
  * Initial document review calls
  * Tool call initiation and completion
  * Chunked document analysis (per chunk)
  * All agent inference operations during redlining workflow
- Thinking content is automatically extracted from various response structures
  and logged in detail with context identifiers for easy tracking
"""

import json
import re
import boto3
from botocore.config import Config
import logging
import time
import sys
import os
from typing import Dict, Any, List
from .tools import retrieve_from_knowledge_base, redline_document, get_tool_definitions, save_analysis_to_dynamodb, parse_conflicts_for_redlining

# Import constants - add parent directories to path
_parent_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
if _parent_dir not in sys.path:
    sys.path.insert(0, _parent_dir)
try:
    import constants
except ImportError:
    # Fallback if constants not available
    class Constants:
        CHUNK_SIZE_CHARACTERS = 30000
        CHUNK_OVERLAP_CHARACTERS = 2000
    constants = Constants()


logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Initialize AWS clients with optimized timeout for Claude Sonnet 4 with thinking
# Optimized timeout based on typical document processing times (2-3 minutes)
# Reference: https://repost.aws/knowledge-center/bedrock-large-model-read-timeouts
bedrock_config = Config(
    read_timeout=300,  # 5 minutes - target completion within 5 minutes
)
bedrock_client = boto3.client('bedrock-runtime', config=bedrock_config)

# Model configuration - Using inference profile for Claude Sonnet 4
CLAUDE_MODEL_ID = "us.anthropic.claude-sonnet-4-20250514-v1:0"
# Fallback: Claude Sonnet 4 with 1M context (same model ID, requires anthropic_beta parameter)
# After 1M fails, we retry the original Sonnet 4
ANTHROPIC_BETA_1M = "context-1m-2025-08-07"  # Beta parameter to enable 1M context window
TEMPERATURE = 1.0  # Must be 1.0 when thinking is enabled
THINKING_BUDGET_TOKENS = 32000  # Increased to 32k for more complex reasoning in document review
MAX_TOKENS = 64000  # Maximum output tokens - must be greater than THINKING_BUDGET_TOKENS per AWS Bedrock requirements

# Graceful queuing configuration for token rate limiting prevention
MAX_RETRIES = 5
BASE_DELAY = 1.0  # Base delay between calls to prevent rate limiting
MAX_DELAY = 6.0  # Max delay for exponential backoff
BACKOFF_MULTIPLIER = 2.0
CALL_SPACING_DELAY = 1.0  # Minimum delay between consecutive calls to prevent token rate limiting

# Global tracking for logging and throttling management
_call_tracker = {
    'total_tool_calls': 0,
    'total_model_calls': 0,
    'total_conflicts_detected': 0,
    'last_call_time': 0
}

def _extract_and_log_thinking(response: Dict[str, Any], context: str = "") -> str:
    """
    Extract thinking content from Claude API response and log it in detail.
    Handles various response structures from AWS Bedrock Converse API.
    
    Args:
        response: Claude API response dictionary
        context: Context string to identify where this thinking occurred (e.g., "initial_review", "chunk_1", "tool_call")
        
    Returns:
        Thinking content as string (empty if not found)
    """
    thinking_content = ""
    
    # Extract thinking from response - check multiple possible locations
    # AWS Bedrock Converse API may structure thinking in different ways
    # Based on actual API responses, thinking is in content blocks as "reasoningContent"
    
    # Method 1: Check for reasoningContent in content blocks - AWS Bedrock Converse API actual location
    # This is where AWS Bedrock returns thinking when enabled via additionalModelRequestFields
    # Check this FIRST since it's the actual location based on API responses
    if response.get("output", {}).get("message", {}).get("content"):
        for idx, content_block in enumerate(response.get("output", {}).get("message", {}).get("content", [])):
            # Check for reasoningContent (the actual field name AWS Bedrock uses)
            if "reasoningContent" in content_block:
                reasoning_val = content_block.get("reasoningContent")
                if isinstance(reasoning_val, str) and reasoning_val:
                    thinking_content = reasoning_val
                    logger.info(f"Found thinking via Method 1: content_block[{idx}].reasoningContent (string, length: {len(thinking_content)})")
                    break
                elif isinstance(reasoning_val, dict):
                    # Handle structured reasoningContent - AWS Bedrock structure: reasoningContent.reasoningText.text
                    reasoning_text_obj = reasoning_val.get("reasoningText")
                    if isinstance(reasoning_text_obj, dict):
                        # reasoningText is a dict with 'text' and 'signature' fields
                        if isinstance(reasoning_text_obj.get("text"), str):
                            thinking_content = reasoning_text_obj.get("text", "")
                            logger.info(f"Found thinking via Method 1: content_block[{idx}].reasoningContent.reasoningText.text (length: {len(thinking_content)})")
                            break
                    elif isinstance(reasoning_text_obj, str):
                        # Fallback: reasoningText might be a string directly
                        thinking_content = reasoning_text_obj
                        logger.info(f"Found thinking via Method 1: content_block[{idx}].reasoningContent.reasoningText (string, length: {len(thinking_content)})")
                        break
                    # Additional fallbacks
                    elif isinstance(reasoning_val.get("text"), str):
                        thinking_content = reasoning_val.get("text", "")
                        logger.info(f"Found thinking via Method 1: content_block[{idx}].reasoningContent.text")
                        break
                    elif isinstance(reasoning_val.get("content"), str):
                        thinking_content = reasoning_val.get("content", "")
                        logger.info(f"Found thinking via Method 1: content_block[{idx}].reasoningContent.content")
                        break
    
    # Method 2: Direct thinking field (string) - fallback location
    if not thinking_content and isinstance(response.get("thinking"), str) and response.get("thinking"):
        thinking_content = response.get("thinking", "")
        logger.info(f"Found thinking via Method 2: direct thinking field")
    
    # Method 3: Thinking in output.thinking (string) - alternative location
    # Check if thinking key exists first (even if None or empty)
    elif not thinking_content and response.get("output"):
        output = response.get("output", {})
        if "thinking" in output:
            thinking_val = output.get("thinking")
            if isinstance(thinking_val, str):
                if thinking_val and thinking_val.strip():  # Only use if not empty or whitespace
                    thinking_content = thinking_val
                    logger.info(f"Found thinking via Method 3: output.thinking (length: {len(thinking_content)})")
                else:  # Empty string means thinking was enabled but not used
                    logger.info(f"DEBUG: output.thinking exists but is empty string (thinking enabled but not used)")
            elif thinking_val is None:
                logger.info(f"DEBUG: output.thinking exists but is None")
            else:
                logger.info(f"DEBUG: output.thinking exists but is unexpected type: {type(thinking_val)}")
    
    # Method 4: Structured thinking object with content/text fields
    elif not thinking_content and isinstance(response.get("thinking"), dict):
        thinking_obj = response.get("thinking", {})
        if isinstance(thinking_obj.get("content"), str):
            thinking_content = thinking_obj.get("content", "")
            logger.info(f"Found thinking via Method 4: thinking.content")
        elif isinstance(thinking_obj.get("text"), str):
            thinking_content = thinking_obj.get("text", "")
            logger.info(f"Found thinking via Method 4: thinking.text")
        elif isinstance(thinking_obj.get("thinking"), str):
            thinking_content = thinking_obj.get("thinking", "")
            logger.info(f"Found thinking via Method 4: thinking.thinking")
    
    # Method 5: Thinking in output.thinking as object
    elif not thinking_content and isinstance(response.get("output", {}).get("thinking"), dict):
        thinking_obj = response.get("output", {}).get("thinking", {})
        if isinstance(thinking_obj.get("content"), str):
            thinking_content = thinking_obj.get("content", "")
            logger.info(f"Found thinking via Method 5: output.thinking.content")
        elif isinstance(thinking_obj.get("text"), str):
            thinking_content = thinking_obj.get("text", "")
            logger.info(f"Found thinking via Method 5: output.thinking.text")
        elif isinstance(thinking_obj.get("thinking"), str):
            thinking_content = thinking_obj.get("thinking", "")
            logger.info(f"Found thinking via Method 5: output.thinking.thinking")
    
    # Method 6: Fallback - Check for thinking field in content blocks (older API versions)
    if not thinking_content and response.get("output", {}).get("message", {}).get("content"):
        for idx, content_block in enumerate(response.get("output", {}).get("message", {}).get("content", [])):
            if content_block.get("thinking"):
                if isinstance(content_block.get("thinking"), str):
                    thinking_content = content_block.get("thinking", "")
                    logger.info(f"Found thinking via Method 6 (fallback): content_block[{idx}].thinking (string)")
                    break
                elif isinstance(content_block.get("thinking"), dict):
                    thinking_obj = content_block.get("thinking", {})
                    if isinstance(thinking_obj.get("content"), str):
                        thinking_content = thinking_obj.get("content", "")
                        logger.info(f"Found thinking via Method 6 (fallback): content_block[{idx}].thinking.content")
                        break
                    elif isinstance(thinking_obj.get("text"), str):
                        thinking_content = thinking_obj.get("text", "")
                        logger.info(f"Found thinking via Method 6 (fallback): content_block[{idx}].thinking.text")
                        break
    
    # Method 6: Check usage metadata for thinking tokens (indicates thinking was used)
    # Note: This doesn't extract thinking content, but confirms thinking was enabled
    if not thinking_content and response.get("usage"):
        usage = response.get("usage", {})
        thinking_tokens = usage.get("thinkingTokens") or usage.get("thinking_tokens") or usage.get("cachedThinkingTokens")
        if thinking_tokens:
            logger.info(f"DEBUG: Found thinking tokens in usage: {thinking_tokens} (but no thinking content found)")
    
    # Log thinking content in detail
    if thinking_content:
        thinking_length = len(thinking_content)
        thinking_preview = thinking_content[:500] if thinking_length > 500 else thinking_content
        
        logger.info(f"=== LLM THINKING [{context}] ===")
        logger.info(f"Thinking content length: {thinking_length} characters")
        logger.info(f"Thinking preview (first 500 chars):\n{thinking_preview}")
        
        # Log full thinking content (split into chunks if too long for single log entry)
        if thinking_length > 5000:
            # Split into chunks for better log readability
            chunk_size = 5000
            num_chunks = (thinking_length // chunk_size) + (1 if thinking_length % chunk_size > 0 else 0)
            logger.info(f"Thinking content is large ({thinking_length} chars), logging in {num_chunks} chunks:")
            
            for i in range(0, thinking_length, chunk_size):
                chunk_num = (i // chunk_size) + 1
                chunk_end = min(i + chunk_size, thinking_length)
                chunk = thinking_content[i:chunk_end]
                logger.info(f"--- Thinking chunk {chunk_num}/{num_chunks} [{context}] ---\n{chunk}")
        else:
            logger.info(f"--- Full thinking content [{context}] ---\n{thinking_content}")
        
        logger.info(f"=== END LLM THINKING [{context}] ===")
    else:
        # Thinking not found - this is normal for tool call responses and continuation responses
        # AWS Bedrock typically only returns thinking in the initial response
        # Only log warning for initial calls, use debug for tool calls/continuations
        if "tool_call" in context or "after_tool" in context:
            logger.debug(f"No thinking content found for context: {context} (normal - thinking typically only in initial response)")
        else:
            logger.debug(f"No thinking content found for context: {context}")
    
    return thinking_content if thinking_content else ""

def _extract_json_only(content: str) -> str:
    """
    Extract only JSON object or array from response content, stripping any explanatory text.
    This ensures we only return valid JSON even if the agent adds explanatory text.
    Prioritizes new format (object with explanation and conflicts) over old format (array).
    
    Args:
        content: Raw response content that may contain JSON plus explanatory text
        
    Returns:
        Clean JSON string (object or array) or empty object/array if no valid JSON found
    """
    if not content:
        return '{"explanation": "", "conflicts": []}'
    
    content_trimmed = content.strip()
    
    # First, try to find JSON object pattern (new format with explanation and conflicts)
    json_object_match = re.search(r'\{[\s\S]*?"conflicts"[\s\S]*?\}', content)
    if json_object_match:
        json_str = json_object_match.group(0)
        # Validate it's actually valid JSON
        try:
            parsed = json.loads(json_str)
            if isinstance(parsed, dict) and "conflicts" in parsed:
                logger.info(f"Extracted valid JSON object from response (length: {len(json_str)} chars)")
                return json_str
        except json.JSONDecodeError:
            logger.warning("Found object-like pattern but not valid JSON, trying bracket matching...")
    
    # Try to find JSON object by bracket matching (for new format)
    start_idx = content_trimmed.find('{')
    if start_idx != -1:
        # Try to parse from { to end, or find matching }
        brace_count = 0
        end_idx = start_idx
        for i in range(start_idx, len(content_trimmed)):
            if content_trimmed[i] == '{':
                brace_count += 1
            elif content_trimmed[i] == '}':
                brace_count -= 1
                if brace_count == 0:
                    end_idx = i + 1
                    break
        
        if brace_count == 0 and end_idx > start_idx:
            json_str = content_trimmed[start_idx:end_idx]
            try:
                parsed = json.loads(json_str)
                if isinstance(parsed, dict):
                    logger.info(f"Extracted JSON object by bracket matching (length: {len(json_str)} chars)")
                    return json_str
            except json.JSONDecodeError:
                pass
    
    # Fallback: try to find JSON array pattern (backwards compatibility)
    json_array_match = re.search(r'\[[\s\S]*?\]', content)
    if json_array_match:
        json_str = json_array_match.group(0)
        # Validate it's actually valid JSON
        try:
            parsed = json.loads(json_str)
            if isinstance(parsed, list):
                logger.info(f"Extracted valid JSON array from response (backwards compatibility, length: {len(json_str)} chars)")
                # Convert array to new format for consistency
                return json.dumps({"explanation": "", "conflicts": parsed})
        except json.JSONDecodeError:
            logger.warning("Found array-like pattern but not valid JSON, trying bracket matching...")
    
    # Try to find JSON array by bracket matching (backwards compatibility)
    start_idx = content_trimmed.find('[')
    if start_idx != -1:
        # Try to parse from [ to end, or find matching ]
        bracket_count = 0
        end_idx = start_idx
        for i in range(start_idx, len(content_trimmed)):
            if content_trimmed[i] == '[':
                bracket_count += 1
            elif content_trimmed[i] == ']':
                bracket_count -= 1
                if bracket_count == 0:
                    end_idx = i + 1
                    break
        
        if bracket_count == 0 and end_idx > start_idx:
            json_str = content_trimmed[start_idx:end_idx]
            try:
                parsed = json.loads(json_str)
                if isinstance(parsed, list):
                    logger.info(f"Extracted JSON array by bracket matching (backwards compatibility, length: {len(json_str)} chars)")
                    # Convert array to new format for consistency
                    return json.dumps({"explanation": "", "conflicts": parsed})
            except json.JSONDecodeError:
                pass
    
    # If all else fails, log warning and return empty object
    logger.warning(f"Could not extract valid JSON from response. Response preview: {content[:200]}...")
    return '{"explanation": "", "conflicts": []}'

def _split_document_into_chunks(doc, chunk_size_characters=30000, chunk_overlap_characters=2000):
    """
    Split a document into chunks using character-based chunking.
    
    Args:
        doc: python-docx Document object
        chunk_size_characters: Number of characters per chunk
        chunk_overlap_characters: Number of characters to overlap between chunks
        
    Returns:
        List of chunk dictionaries with bytes, chunk_num, start_char, end_char
    """
    from docx import Document
    import io
    
    # Import constants
    try:
        import constants
        chunk_size = getattr(constants, 'CHUNK_SIZE_CHARACTERS', chunk_size_characters)
        overlap = getattr(constants, 'CHUNK_OVERLAP_CHARACTERS', chunk_overlap_characters)
    except (ImportError, AttributeError):
        # Fallback to function parameters
        chunk_size = chunk_size_characters
        overlap = chunk_overlap_characters
    
    chunks = []
    
    # For DOCX: Extract full text and chunk by characters
    # Extract all text from document
    full_text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text_parts.append(para.text)
    
    full_text = '\n\n'.join(full_text_parts)
    total_chars = len(full_text)
    
    if total_chars == 0:
        # Empty document - return single empty chunk
        buffer = io.BytesIO()
        doc.save(buffer)
        chunks.append({
            'bytes': buffer.getvalue(),
            'chunk_num': 0,
            'start_char': 0,
            'end_char': 0
        })
        return chunks
    
    start_char = 0
    chunk_num = 0
    
    while start_char < total_chars:
        end_char = min(start_char + chunk_size, total_chars)
        
        # Extract chunk text
        chunk_text = full_text[start_char:end_char]
        
        # Create a new document for this chunk
        chunk_doc = Document()
        
        # Split chunk text into paragraphs and add to document
        for para_text in chunk_text.split('\n\n'):
            if para_text.strip():
                chunk_doc.add_paragraph(para_text.strip())
        
        # Save to bytes
        buffer = io.BytesIO()
        chunk_doc.save(buffer)
        chunk_bytes = buffer.getvalue()
        
        chunks.append({
            'bytes': chunk_bytes,
            'chunk_num': chunk_num,
            'start_char': start_char,
            'end_char': end_char
        })
        
        chunk_num += 1
        # Move start position with overlap
        start_char = end_char - overlap
        # Break when we've processed all content (end_char reached total_chars)
        if end_char >= total_chars:
            break
    
    return chunks



def get_kb_id_by_name(name: str) -> str:
    """Resolve Knowledge Base ID from Name."""
    try:
        client = boto3.client('bedrock-agent')
        paginator = client.get_paginator('list_knowledge_bases')
        for page in paginator.paginate(MaxResults=100):
            for kb in page.get('knowledgeBaseSummaries', []):
                if kb.get('name') == name:
                    logger.info(f"Resolved Knowledge Base ID {kb.get('knowledgeBaseId')} for name {name}")
                    return kb.get('knowledgeBaseId')
        logger.warning(f"Knowledge Base with name {name} not found")
    except Exception as e:
        logger.error(f"Error resolving KB ID for name {name}: {e}")
    return None

class Model:
    """
    Handles document review using Claude 4 Sonnet thinking with tool calling.
    """
    
    def __init__(self, knowledge_base_id: str, region: str):
        self.knowledge_base_id = knowledge_base_id
        self.region = region
        
        # Resolve Knowledge Base ID if it is missing or a placeholder
        if (not self.knowledge_base_id or self.knowledge_base_id == "placeholder") and os.environ.get('KNOWLEDGE_BASE_NAME'):
            resolved_id = get_kb_id_by_name(os.environ.get('KNOWLEDGE_BASE_NAME'))
            if resolved_id:
                self.knowledge_base_id = resolved_id
                logger.info(f"Model initialized with resolved KB ID: {self.knowledge_base_id}")
            else:
                logger.warning(f"Failed to resolve KB ID from name: {os.environ.get('KNOWLEDGE_BASE_NAME')}")
        
        self.tools = get_tool_definitions()
    

    
    def _get_bucket_name(self, bucket_type: str) -> str:
        """Get the appropriate bucket name based on type."""
        import os
        if bucket_type == "knowledge":
            return os.environ.get("KNOWLEDGE_BUCKET")
        elif bucket_type == "user_documents":
            return os.environ.get("USER_DOCUMENTS_BUCKET")
        elif bucket_type == "agent_processing":
            return os.environ.get("AGENT_PROCESSING_BUCKET")
        else:
            raise ValueError(f"Invalid bucket_type: {bucket_type}")
    
    def _get_document_format(self, document_s3_key: str) -> str:
        """Determine document format based on file extension for Converse API."""
        file_extension = document_s3_key.lower().split('.')[-1]
        
        # Converse API supported formats
        supported_formats = ['csv', 'doc', 'docx', 'xls', 'xlsx', 'html', 'txt', 'md']
        
        if file_extension in supported_formats:
            return file_extension
        else:
            # Default to DOCX if extension not recognized
            logger.warning(f"Unknown file extension '{file_extension}', defaulting to DOCX format")
            return 'docx'
    
    def _sanitize_filename_for_converse(self, filename: str) -> str:
        """
        Sanitize filename to meet Converse API requirements:
        - Only alphanumeric, whitespace, hyphens, parentheses, square brackets
        - No more than one consecutive whitespace character
        """
        import re
        
        # If filename contains UUID or timestamp patterns, extract just the original name
        # Pattern: uuid_originalname.ext or timestamp/uuid_originalname.ext
        if '_' in filename:
            # Try to extract original filename after last underscore
            parts = filename.split('_')
            if len(parts) > 1:
                # Keep everything after the last underscore (likely the original filename)
                filename = parts[-1]
        
        # Remove any invalid characters - keep only allowed ones
        # Allowed: alphanumeric, whitespace, hyphens, parentheses, square brackets
        sanitized = re.sub(r'[^a-zA-Z0-9\s\-\(\)\[\]]', '', filename)
        
        # Replace multiple consecutive whitespace with single space
        sanitized = re.sub(r'\s+', ' ', sanitized)
        
        # Trim whitespace from start/end
        sanitized = sanitized.strip()
        
        # If sanitization left us with empty string, use default
        if not sanitized:
            sanitized = "vendor-submission"
        
        logger.info(f"Sanitized filename from '{filename}' to '{sanitized}'")
        return sanitized
    
    def _call_claude_without_tools(self, messages: List[Dict[str, Any]], retry_count: int = 0, use_1m_context: bool = False, tried_1m: bool = False, enable_thinking: bool = True, temperature: float = None) -> Dict[str, Any]:
        """
        Call Claude without tool support using Converse API.
        Use this when KB results are already pre-loaded in the prompt.
        Implements graceful queuing with call spacing to prevent token rate limiting.
        Supports fallback: Primary Sonnet 4 (once) -> Sonnet 4 1M -> Retry Sonnet 4 with backoff
        
        Args:
            messages: List of message dictionaries
            retry_count: Current retry attempt number
            use_1m_context: Whether to use 1M context version (fallback)
            tried_1m: Whether we've already attempted 1M context (prevents loops)
            enable_thinking: Whether to enable thinking mode (default: True). If False, allows lower temperature.
            temperature: Temperature to use (default: TEMPERATURE constant). If thinking disabled, can use 0.0-0.3.
        """
        
        # Implement graceful call spacing to prevent token rate limiting
        current_time = time.time()
        time_since_last_call = current_time - _call_tracker['last_call_time']
        
        if time_since_last_call < CALL_SPACING_DELAY:
            wait_time = CALL_SPACING_DELAY - time_since_last_call
            logger.info(f"Graceful queuing: waiting {wait_time:.2f}s to prevent token rate limiting")
            time.sleep(wait_time)
        
        # Always use Claude Sonnet 4 - just toggle 1M context via beta parameter
        current_model_id = CLAUDE_MODEL_ID
        
        # Log attempt (but don't increment counter until successful)
        model_name = "Sonnet 4 1M" if use_1m_context else "Sonnet 4"
        logger.info(f"Calling Claude ({model_name}: {current_model_id}) with {len(messages)} messages without tools (attempt {retry_count + 1}, use_1m_context={use_1m_context}) - Total successful calls so far: {_call_tracker['total_model_calls']}")
        
        try:
            # Determine temperature: use provided value, or TEMPERATURE constant, or 0.0 if thinking disabled
            if temperature is None:
                if enable_thinking:
                    temperature = TEMPERATURE  # Must be 1.0 when thinking enabled
                else:
                    temperature = 0.0  # Deterministic when thinking disabled
            
            # Prepare inference config
            inference_config = {
                "temperature": temperature,
                "maxTokens": MAX_TOKENS
            }
            
            # Prepare API call parameters - NO tools
            # Step functions use their own specialized prompts in user messages, so no system prompt needed
            api_params = {
                "modelId": current_model_id,
                "messages": messages,
                "inferenceConfig": inference_config
                # NO toolConfig - tools disabled
            }
            
            # Add thinking config only if enabled
            if enable_thinking:
                api_params["additionalModelRequestFields"] = {
                    "thinking": {
                        "type": "enabled",
                        "budget_tokens": THINKING_BUDGET_TOKENS
                    }
                }
            
            # Add 1M context beta parameter if using 1M fallback
            # AWS Bedrock expects anthropic_beta to be a list of strings
            if use_1m_context:
                api_params["additionalModelRequestFields"]["anthropic_beta"] = [ANTHROPIC_BETA_1M]
            
            # Call Bedrock using Converse API (supports document attachments)
            response = bedrock_client.converse(**api_params)
            
            # SUCCESS: Only now increment the counter for successful calls
            _call_tracker['total_model_calls'] += 1
            _call_tracker['last_call_time'] = time.time()
            
            logger.info(f"Claude API call successful! Total successful model calls: {_call_tracker['total_model_calls']}")
            
            # Extract and log thinking content
            thinking_context = f"inference_call_{_call_tracker['total_model_calls']}"
            _extract_and_log_thinking(response, thinking_context)
            
            # No tool calls possible - return response directly
            return response
            
        except Exception as e:
            error_msg = str(e)
            error_type = type(e).__name__
            
            # Check for throttling errors and implement exponential backoff
            is_throttling = ("ThrottlingException" in error_msg or "Too many tokens" in error_msg or 
                            "rate" in error_msg.lower() or "throttl" in error_msg.lower() or
                            "limit" in error_msg.lower())
            
            # Check for transient errors that should be retried
            is_transient = (
                "ServiceUnavailableException" in error_type or
                "InternalServerError" in error_type or
                "InternalFailure" in error_msg or
                "ServiceUnavailable" in error_msg or
                "timeout" in error_msg.lower() or
                "Timeout" in error_type or
                "Connection" in error_type or
                "ReadTimeout" in error_type or
                "502" in error_msg or  # Bad Gateway
                "503" in error_msg or  # Service Unavailable
                "504" in error_msg     # Gateway Timeout
            )
            
            # Handle throttling with exponential backoff
            if is_throttling and retry_count < MAX_RETRIES:
                wait_time = min(2 ** retry_count, MAX_BACKOFF_SECONDS)
                logger.warning(f"Claude API throttling error detected ({error_type}), retrying in {wait_time} seconds (attempt {retry_count + 1}/{MAX_RETRIES})")
                time.sleep(wait_time)
                return self._call_claude_without_tools(messages, retry_count + 1, use_1m_context, tried_1m)
            
            # Handle transient errors
            if is_transient and retry_count < MAX_RETRIES:
                wait_time = min(2 ** retry_count, MAX_BACKOFF_SECONDS)
                logger.warning(f"Claude API transient error detected ({error_type}), retrying in {wait_time} seconds (attempt {retry_count + 1}/{MAX_RETRIES})")
                time.sleep(wait_time)
                return self._call_claude_without_tools(messages, retry_count + 1, use_1m_context, tried_1m)
            
            # Try 1M context fallback if primary model fails and we haven't tried it yet
            if not tried_1m and not use_1m_context:
                logger.warning(f"Sonnet 4 failed on first attempt. Attempting fallback to Sonnet 4 1M")
                return self._call_claude_without_tools(messages, retry_count, use_1m_context=True, tried_1m=True)
            
            # If 1M context also failed, retry with exponential backoff
            if use_1m_context and retry_count < MAX_RETRIES:
                wait_time = min(2 ** retry_count, MAX_BACKOFF_SECONDS)
                logger.warning(f"Sonnet 4 1M also failed. Retrying Sonnet 4 with exponential backoff")
                logger.warning(f"Retrying Sonnet 4 in {wait_time} seconds (attempt {retry_count + 1})")
                time.sleep(wait_time)
                return self._call_claude_without_tools(messages, retry_count + 1, use_1m_context=False, tried_1m=True)
            
            # Max retries exceeded
            if retry_count >= MAX_RETRIES:
                logger.error(f"Max retries exceeded for Claude API error after {MAX_RETRIES} attempts")
            
            # Re-raise the exception if we can't handle it
            raise
    
    def _call_claude_with_tools(self, messages: List[Dict[str, Any]], retry_count: int = 0, use_1m_context: bool = False, tried_1m: bool = False, enable_thinking: bool = True, temperature: float = None) -> Dict[str, Any]:
        """
        Call Claude with tool support using Converse API.
        Implements graceful queuing with call spacing to prevent token rate limiting.
        Supports fallback: Primary Sonnet 4 (once) -> Sonnet 4 1M -> Retry Sonnet 4 with backoff
        
        Args:
            messages: List of message dictionaries
            retry_count: Current retry attempt number
            use_1m_context: Whether to use 1M context version (fallback)
            tried_1m: Whether we've already attempted 1M context (prevents loops)
            enable_thinking: Whether to enable thinking mode (default: True). If False, allows lower temperature.
            temperature: Temperature to use (default: TEMPERATURE constant). If thinking disabled, can use 0.0-0.3.
        """
        
        # Implement graceful call spacing to prevent token rate limiting
        current_time = time.time()
        time_since_last_call = current_time - _call_tracker['last_call_time']
        
        if time_since_last_call < CALL_SPACING_DELAY:
            wait_time = CALL_SPACING_DELAY - time_since_last_call
            logger.info(f"Graceful queuing: waiting {wait_time:.2f}s to prevent token rate limiting")
            time.sleep(wait_time)
        
        # Always use Claude Sonnet 4 - just toggle 1M context via beta parameter
        current_model_id = CLAUDE_MODEL_ID
        
        # Log attempt (but don't increment counter until successful)
        model_name = "Sonnet 4 1M" if use_1m_context else "Sonnet 4"
        logger.info(f"Calling Claude ({model_name}: {current_model_id}) with {len(messages)} messages and {len(self.tools)} tools (attempt {retry_count + 1}, use_1m_context={use_1m_context}) - Total successful calls so far: {_call_tracker['total_model_calls']}")
        
        try:
            # Determine temperature: use provided value, or TEMPERATURE constant, or 0.0 if thinking disabled
            if temperature is None:
                if enable_thinking:
                    temperature = TEMPERATURE  # Must be 1.0 when thinking enabled
                else:
                    temperature = 0.0  # Deterministic when thinking disabled
            
            # Prepare inference config
            inference_config = {
                "temperature": temperature,
                "maxTokens": MAX_TOKENS
            }
            
            # Prepare API call parameters
            # Step functions use their own specialized prompts in user messages, so no system prompt needed
            api_params = {
                "modelId": current_model_id,
                "messages": messages,
                "inferenceConfig": inference_config,
                "toolConfig": {"tools": self.tools}
            }
            
            # Add thinking config only if enabled
            if enable_thinking:
                api_params["additionalModelRequestFields"] = {
                    "thinking": {
                        "type": "enabled",
                        "budget_tokens": THINKING_BUDGET_TOKENS
                    }
                }
            
            # Add 1M context beta parameter if using 1M fallback
            # AWS Bedrock expects anthropic_beta to be a list of strings
            if use_1m_context:
                api_params["additionalModelRequestFields"]["anthropic_beta"] = [ANTHROPIC_BETA_1M]
            
            # Call Bedrock using Converse API (supports document attachments)
            response = bedrock_client.converse(**api_params)
            
            # SUCCESS: Only now increment the counter for successful calls
            _call_tracker['total_model_calls'] += 1
            _call_tracker['last_call_time'] = time.time()
            
            logger.info(f"Claude API call successful! Total successful model calls: {_call_tracker['total_model_calls']}")
            
            # Extract and log thinking content (only the reasoning text will be logged)
            thinking_context = f"inference_call_{_call_tracker['total_model_calls']}"
            if response.get("stopReason") == "tool_use":
                thinking_context += "_before_tool_use"
            _extract_and_log_thinking(response, thinking_context)
            
            # Handle tool calls if present
            if response.get("stopReason") == "tool_use":
                return self._handle_tool_calls(messages, response, enable_thinking=enable_thinking, temperature=temperature)
            
            return response
            
        except Exception as e:
            error_msg = str(e)
            error_type = type(e).__name__
            
            # Check for throttling errors and implement exponential backoff
            is_throttling = ("ThrottlingException" in error_msg or "Too many tokens" in error_msg or 
                            "rate" in error_msg.lower() or "throttl" in error_msg.lower() or
                            "limit" in error_msg.lower())
            
            # Check for transient errors that should be retried
            is_transient = (
                "ServiceUnavailableException" in error_type or
                "InternalServerError" in error_type or
                "InternalFailure" in error_msg or
                "ServiceUnavailable" in error_msg or
                "timeout" in error_msg.lower() or
                "Timeout" in error_type or
                "Connection" in error_type or
                "ReadTimeout" in error_type or
                "502" in error_msg or  # Bad Gateway
                "503" in error_msg or  # Service Unavailable
                "504" in error_msg     # Gateway Timeout
            )
            
            # Check for validation errors that should NOT be retried (configuration issues)
            is_validation_error = (
                "ValidationException" in error_type or
                "ValidationError" in error_type or
                "InvalidParameter" in error_type or
                "InvalidRequest" in error_msg
            )
            
            # Retry logic for throttling and transient errors
            if (is_throttling or is_transient) and not is_validation_error:
                # On first failure (retry_count=0), try Sonnet 4 1M if not already tried
                if retry_count == 0 and not use_1m_context and not tried_1m:
                    logger.warning(f"Sonnet 4 failed on first attempt. Attempting fallback to Sonnet 4 1M")
                    try:
                        return self._call_claude_with_tools(messages, retry_count=0, use_1m_context=True, tried_1m=True)
                    except Exception as fallback_1m_error:
                        # If 1M fails, go back to Sonnet 4 and continue retrying with backoff
                        logger.warning(f"Sonnet 4 1M also failed. Retrying Sonnet 4 with exponential backoff")
                        delay = min(BASE_DELAY * (BACKOFF_MULTIPLIER ** retry_count), MAX_DELAY)
                        error_category = "throttling" if is_throttling else "transient"
                        logger.warning(f"Retrying Sonnet 4 in {delay} seconds (attempt {retry_count + 2})")
                        time.sleep(delay)
                        return self._call_claude_with_tools(messages, retry_count + 1, use_1m_context=False, tried_1m=True)
                
                # Continue retrying Sonnet 4 with exponential backoff
                if retry_count < MAX_RETRIES:
                    delay = min(BASE_DELAY * (BACKOFF_MULTIPLIER ** retry_count), MAX_DELAY)
                    error_category = "throttling" if is_throttling else "transient"
                    logger.warning(f"Claude API {error_category} error detected ({error_type}), retrying in {delay} seconds (attempt {retry_count + 1}/{MAX_RETRIES + 1})")
                    time.sleep(delay)
                    return self._call_claude_with_tools(messages, retry_count + 1, use_1m_context=False, tried_1m=tried_1m)
                else:
                    # Max retries exceeded
                    error_category = "throttling" if is_throttling else "transient"
                    logger.error(f"Max retries exceeded for Claude API {error_category} error after {MAX_RETRIES + 1} attempts")
                    raise Exception(f"Claude API {error_category} error limit exceeded after {MAX_RETRIES + 1} retries: {str(e)}")
            else:
                # For validation errors, try Sonnet 4 1M if not already tried, then retry Sonnet 4
                if not use_1m_context and is_validation_error and not tried_1m:
                    logger.warning(f"Validation error on Sonnet 4. Attempting fallback to Sonnet 4 1M")
                    try:
                        return self._call_claude_with_tools(messages, retry_count=0, use_1m_context=True, tried_1m=True)
                    except Exception as fallback_1m_error:
                        # If 1M also fails, retry original Sonnet 4
                        logger.warning(f"Sonnet 4 1M also failed with validation error. Retrying original Sonnet 4")
                        try:
                            return self._call_claude_with_tools(messages, retry_count=0, use_1m_context=False, tried_1m=True)
                        except Exception as final_error:
                            logger.error(f"All attempts failed with validation errors. Sonnet 4: {str(e)}, Sonnet 4 1M: {str(fallback_1m_error)}, Sonnet 4 retry: {str(final_error)}")
                            raise Exception(f"Claude API validation error on all attempts. Sonnet 4: {str(e)}, Sonnet 4 1M: {str(fallback_1m_error)}, Sonnet 4 retry: {str(final_error)}")
                
                # Non-retryable error - don't retry
                logger.error(f"Error calling Claude (non-retryable {error_type}): {str(e)}")
                raise
    
    def _handle_tool_calls(self, messages: List[Dict[str, Any]], claude_response: Dict[str, Any], enable_thinking: bool = True, temperature: float = None) -> Dict[str, Any]:
        """
        Handle tool calls from Claude and continue the conversation.
        
        Args:
            messages: Current conversation messages
            claude_response: Claude's response containing tool calls
            enable_thinking: Whether thinking mode is enabled (passed through from original call)
            temperature: Temperature setting (passed through from original call)
        """
        
        # Log thinking from the initial tool use response
        logger.info("=== TOOL CALL DETECTED - Logging thinking before tool execution ===")
        _extract_and_log_thinking(claude_response, f"tool_call_initiation_{_call_tracker['total_tool_calls'] + 1}")
        
        # Add Claude's response to messages (Converse API format)
        messages.append({
            "role": "assistant",
            "content": claude_response["output"]["message"]["content"]
        })
        
        tool_results = []
        
        # Process each tool call (Converse API format)
        for content_block in claude_response["output"]["message"]["content"]:
            if content_block.get("toolUse"):
                tool_use = content_block["toolUse"]
                tool_name = tool_use["name"]
                tool_input = tool_use["input"]
                tool_use_id = tool_use["toolUseId"]
                
                # Update tool call tracking (safe to increment here - no retries for tools)
                _call_tracker['total_tool_calls'] += 1
                
                logger.info(f"Executing tool: {tool_name} with input: {tool_input} - Total tool calls: {_call_tracker['total_tool_calls']}")
                
                # Execute the tool using existing tools.py functions
                try:
                    if tool_name == "retrieve_from_knowledge_base":
                        result = retrieve_from_knowledge_base(
                            query=tool_input["query"],
                            max_results=tool_input.get("max_results", 100),
                            knowledge_base_id=self.knowledge_base_id,
                            region=self.region
                        )
                        logger.info(f"Tool {tool_name} executed successfully")
                    else:
                        result = {"error": f"Unknown tool: {tool_name}"}
                        logger.warning(f"Unknown tool requested: {tool_name}")
                    
                    tool_results.append({
                        "tool_name": tool_name,
                        "input": tool_input,
                        "result": result
                    })
                    
                    # Add tool result to messages (Converse API format)
                    messages.append({
                        "role": "user",
                        "content": [
                            {
                                "toolResult": {
                                    "toolUseId": tool_use_id,
                                    "content": [{"json": result}]
                                }
                            }
                        ]
                    })
                    
                except Exception as e:
                    logger.error(f"Error executing tool {tool_name}: {str(e)}")
                    error_result = {"error": str(e)}
                    tool_results.append({
                        "tool_name": tool_name,
                        "input": tool_input,
                        "result": error_result
                    })
                    
                    # Add error to messages (Converse API format)
                    messages.append({
                        "role": "user",
                        "content": [
                            {
                                "toolResult": {
                                    "toolUseId": tool_use_id,
                                    "content": [{"json": error_result}]
                                }
                            }
                        ]
                    })
        
        # Continue the conversation with tool results
        # CRITICAL: Pass through enable_thinking and temperature to maintain consistency
        # If thinking was disabled initially, keep it disabled to avoid validation errors
        logger.info("=== CONTINUING CONVERSATION AFTER TOOL EXECUTION ===")
        final_response = self._call_claude_with_tools(messages, enable_thinking=enable_thinking, temperature=temperature)
        
        # Log thinking from the final response after tool execution
        logger.info("=== LOGGING THINKING AFTER TOOL EXECUTION ===")
        _extract_and_log_thinking(final_response, f"after_tool_execution_{_call_tracker['total_tool_calls']}")
        
        final_response["tool_results"] = tool_results
        
        return final_response
 
