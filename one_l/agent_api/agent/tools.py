"""
Tool functions for document review agent.
Provides knowledge base retrieval and document red-lining capabilities.
"""

import json
import boto3
from botocore.config import Config
import os
import logging
import re
import time
import hashlib
import unicodedata
from typing import Dict, Any, List, Optional
from collections import defaultdict
from docx import Document
from docx.shared import RGBColor, Pt, Inches
import io

# Pydantic models for output validation
try:
    from pydantic import BaseModel, Field, field_validator, ConfigDict
    PYDANTIC_AVAILABLE = True
except ImportError:
    PYDANTIC_AVAILABLE = False
    BaseModel = None
    Field = None
    field_validator = None
    ConfigDict = None


logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Pydantic model for conflict validation
if PYDANTIC_AVAILABLE:
    class ConflictModel(BaseModel):
        """Pydantic model for validating conflict data from LLM output."""
        clarification_id: str = Field(..., description="Vendor's ID or Additional-[#] for other findings")
        vendor_quote: str = Field(..., description="Exact text verbatim OR 'N/A - Missing provision' for omissions")
        summary: str = Field(..., description="20-40 word context")
        source_doc: str = Field(..., description="Name of the actual Massachusetts source document retrieved from the knowledge base, OR 'N/A – General risk language not tied to a specific Massachusetts clause'")
        clause_ref: str = Field(default="N/A", description="Specific section or 'N/A' if not applicable")
        conflict_type: str = Field(..., description="adds/deletes/modifies/contradicts/omits required/reverses obligation")
        rationale: str = Field(..., description="≤50 words on legal impact")
        
        @field_validator('source_doc')
        @classmethod
        def validate_source_doc(cls, v: str) -> str:
            """Normalize source_doc but allow any value - no hard restrictions on conflict removal."""
            v_str = str(v).strip()
            # Return as-is, no validation/rejection - allow all conflicts through
            return v_str if v_str else ""
        
        @field_validator('vendor_quote')
        @classmethod
        def validate_vendor_quote(cls, v: str) -> str:
            """Clean up vendor quote by removing surrounding quotes."""
            v_str = str(v).strip()
            if v_str.startswith('"') and v_str.endswith('"'):
                v_str = v_str[1:-1]
            if not v_str or v_str.lower() in ['n/a', 'na', 'none', 'n.a.', 'n.a', 'not available'] or len(v_str) < 5:
                raise ValueError(f"vendor_quote must be meaningful text (at least 5 chars), got: '{v_str}'")
            return v_str
        
        model_config = ConfigDict(
            extra='forbid',  # Reject extra fields
            str_strip_whitespace=True  # Auto-strip strings
        )
    
    class RedliningResponseModel(BaseModel):
        """Pydantic model for validating the redlining response structure with explanation and conflicts."""
        explanation: str = Field(..., description="Justification/explanation in the form of text so the model can give more context")
        conflicts: List[ConflictModel] = Field(default_factory=list, description="Array of conflicts")
        
        model_config = ConfigDict(
            extra='forbid',  # Reject extra fields
            str_strip_whitespace=True  # Auto-strip strings
        )

# Initialize AWS clients with optimized timeout for knowledge base retrieval
# Knowledge base queries are typically fast (under 30 seconds)
# Reference: https://repost.aws/knowledge-center/bedrock-large-model-read-timeouts
bedrock_agent_config = Config(
    read_timeout=120,  # 2 minutes - more than sufficient for knowledge base queries
)
bedrock_agent_client = boto3.client('bedrock-agent-runtime', config=bedrock_agent_config)
s3_client = boto3.client('s3')

# Knowledge base optimization constants - TUNED FOR MAXIMUM CONFLICT DETECTION
MAX_CHUNK_SIZE = 3000  # Increased tokens per chunk for more context
MIN_RELEVANCE_SCORE = 0.5  # Lowered threshold to capture more potentially relevant content
OPTIMAL_RESULTS_PER_QUERY = 50  # Increased results per query for comprehensive coverage
DEDUPLICATION_THRESHOLD = 0.90  # Slightly higher threshold to allow more similar content variations

# Exponential backoff configuration for throttling resilience
MAX_RETRIES = 5
BASE_DELAY = 1.0
MAX_DELAY = 32.0
BACKOFF_MULTIPLIER = 2.0

# Global cache for session-based deduplication
_content_cache = {}
_query_cache = {}

def _calculate_content_signature(text: str) -> str:
    """Calculate semantic signature for deduplication."""
    # Normalize text for consistent comparison
    normalized = re.sub(r'\s+', ' ', text.lower().strip())
    normalized = re.sub(r'[^\w\s]', '', normalized)
    return hashlib.sha256(normalized.encode()).hexdigest()[:16]

def _is_duplicate_content(text: str, threshold: float = DEDUPLICATION_THRESHOLD) -> bool:
    """Check if content is duplicate using semantic similarity."""
    signature = _calculate_content_signature(text)
    
    if signature in _content_cache:
        return True
    
    # Simple similarity check with existing content signatures
    # Create a list of keys to iterate over to avoid "dictionary changed size during iteration" error
    # when multiple threads access the cache concurrently
    existing_signatures = list(_content_cache.keys())
    for existing_sig in existing_signatures:
        matches = sum(c1 == c2 for c1, c2 in zip(signature, existing_sig) if len(signature) == len(existing_sig))
        if len(existing_sig) > 0 and matches / len(existing_sig) > threshold:
            return True
    
    _content_cache[signature] = text
    return False

def _chunk_content_intelligently(text: str, max_size: int = MAX_CHUNK_SIZE) -> List[str]:
    """Intelligently chunk content preserving semantic boundaries."""
    if len(text) <= max_size * 4:  # Approximate token conversion (4 chars = 1 token)
        return [text]
    
    # Split on semantic boundaries (sentences, then paragraphs)
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= max_size * 4:
            current_chunk += sentence + " "
        else:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            current_chunk = sentence + " "
    
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks

def _filter_and_prioritize_results(results: List[Dict], max_results: int, terms_profile: str = None) -> List[Dict]:
    """Filter results by relevance and prioritize for optimal context usage.
    
    Args:
        results: List of retrieval results
        max_results: Maximum number of results to return
        terms_profile: Optional terms profile ('general_terms', 'it_terms_updated', 'it_terms_old') for filtering
    """
    import os
    
    # Filter by minimum relevance score
    filtered_results = [
        r for r in results 
        if r.get('score', 0) >= MIN_RELEVANCE_SCORE
    ]
    
    # Map terms_profile to bucket name patterns (defined early for use in sorting)
    terms_bucket_patterns = {
        'general_terms': [
            'general-terms',  # Bucket name segment
            'general_terms',  # Alternative format
            'onel-prod-general-terms',  # Full bucket name (stack prefix)
            'form_commonwealth-terms-and-conditions'  # Document filename pattern
        ],
        'it_terms_updated': [
            'it-terms-updated',  # Bucket name segment
            'it_terms_updated',  # Alternative format
            'onel-prod-it-terms-updated',  # Full bucket name (stack prefix)
            'Updated IT Terms'  # Document filename pattern
        ],
        'it_terms_old': [
            'it-terms-old',  # Bucket name segment
            'it_terms_old',  # Alternative format
            'onel-prod-it-terms-old'  # Full bucket name (stack prefix)
        ]
    }
    
    # Filter by terms profile if specified
    if not terms_profile:
        logger.debug(f"TERMS_FILTER: No terms_profile provided, skipping filtering (terms_profile={terms_profile})")
    
    if terms_profile:
        logger.info(f"TERMS_FILTER_START: Applying filter for terms_profile={terms_profile}, total_results={len(filtered_results)}")
        
        # Get allowed patterns for selected terms profile
        allowed_patterns = terms_bucket_patterns.get(terms_profile, [])
        
        # Get excluded patterns (other terms buckets)
        excluded_patterns = []
        for profile, patterns in terms_bucket_patterns.items():
            if profile != terms_profile:
                excluded_patterns.extend(patterns)
        
        logger.info(f"TERMS_FILTER_INFO: Selected profile '{terms_profile}' allows patterns: {allowed_patterns}")
        logger.info(f"TERMS_FILTER_INFO: Excluding patterns from other profiles: {excluded_patterns}")
        
        # Log document sources before filtering for debugging
        if filtered_results:
            source_buckets = {}
            for r in filtered_results[:10]:  # Log first 10 for brevity
                location = r.get('location', {})
                s3_location = location.get('s3Location', {})
                s3_uri = s3_location.get('uri', '')
                source = r.get('source', '')
                bucket_name = 'unknown'
                if s3_uri:
                    # Extract bucket name from S3 URI (s3://bucket-name/...)
                    if 's3://' in s3_uri:
                        bucket_name = s3_uri.split('/')[2] if len(s3_uri.split('/')) > 2 else 'unknown'
                elif 'general-terms' in source.lower() or 'general_terms' in source.lower():
                    bucket_name = 'general-terms'
                elif 'it-terms-updated' in source.lower() or 'it_terms_updated' in source.lower():
                    bucket_name = 'it-terms-updated'
                elif 'it-terms-old' in source.lower() or 'it_terms_old' in source.lower():
                    bucket_name = 'it-terms-old'
                
                source_buckets[bucket_name] = source_buckets.get(bucket_name, 0) + 1
            
            logger.info(f"TERMS_FILTER_INFO: Document sources before filtering (sample): {source_buckets}")
        
        # Filter out documents from non-selected terms buckets
        # Use same priority-based matching as boost logic to avoid false exclusions
        if excluded_patterns:
            filtered_by_terms = []
            excluded_count = 0
            excluded_sources = {}
            included_sources = {}
            
            for r in filtered_results:
                location = r.get('location', {})
                s3_location = location.get('s3Location', {})
                s3_uri = s3_location.get('uri', '').lower()
                s3_key = s3_location.get('key', '').lower()
                
                # Extract bucket name from S3 URI (format: s3://bucket-name/path/to/file.pdf)
                bucket_name_from_uri = ''
                if s3_uri and s3_uri.startswith('s3://'):
                    try:
                        bucket_name_from_uri = s3_uri.split('/')[2] if len(s3_uri.split('/')) > 2 else ''
                    except:
                        bucket_name_from_uri = ''
                
                # Check if this document belongs to an excluded terms bucket
                # Priority 1: Check S3 bucket name (most reliable)
                is_excluded = False
                if bucket_name_from_uri:
                    is_excluded = any(pattern.lower() in bucket_name_from_uri for pattern in excluded_patterns)
                
                # Priority 2: Check S3 key path if bucket didn't match
                if not is_excluded and s3_key:
                    is_excluded = any(pattern.lower() in s3_key for pattern in excluded_patterns)
                
                # Priority 3: Check specific filename patterns (only for exact matches)
                if not is_excluded:
                    source = r.get('source', '').lower()
                    filename_lower = os.path.basename(source) if source else ''
                    specific_filename_patterns = [
                        'form_commonwealth-terms-and-conditions',
                        'updated it terms'
                    ]
                    for pattern in excluded_patterns:
                        if pattern.lower() in specific_filename_patterns:
                            if pattern.lower() in filename_lower:
                                is_excluded = True
                                break
                
                if not is_excluded:
                    filtered_by_terms.append(r)
                    # Track included sources using bucket name
                    bucket_key = 'unknown'
                    if 'general-terms' in bucket_name_from_uri or 'general_terms' in bucket_name_from_uri:
                        bucket_key = 'general-terms'
                    elif 'it-terms-updated' in bucket_name_from_uri or 'it_terms_updated' in bucket_name_from_uri:
                        bucket_key = 'it-terms-updated'
                    elif 'it-terms-old' in bucket_name_from_uri or 'it_terms_old' in bucket_name_from_uri:
                        bucket_key = 'it-terms-old'
                    elif bucket_name_from_uri:
                        bucket_key = bucket_name_from_uri  # Use actual bucket name if not a terms bucket
                    included_sources[bucket_key] = included_sources.get(bucket_key, 0) + 1
                else:
                    excluded_count += 1
                    # Track excluded sources using bucket name
                    bucket_key = 'unknown'
                    if 'general-terms' in bucket_name_from_uri or 'general_terms' in bucket_name_from_uri:
                        bucket_key = 'general-terms'
                    elif 'it-terms-updated' in bucket_name_from_uri or 'it_terms_updated' in bucket_name_from_uri:
                        bucket_key = 'it-terms-updated'
                    elif 'it-terms-old' in bucket_name_from_uri or 'it_terms_old' in bucket_name_from_uri:
                        bucket_key = 'it-terms-old'
                    excluded_sources[bucket_key] = excluded_sources.get(bucket_key, 0) + 1
                    source = r.get('source', '')
                    logger.debug(f"TERMS_FILTER_EXCLUDE: Excluding document - source: {source[:80]}, bucket: {bucket_name_from_uri}, s3_uri: {s3_uri[:80] if s3_uri else 'N/A'}")
            
            logger.info(f"TERMS_FILTER_RESULT: Excluded {excluded_count} documents from non-selected terms buckets")
            logger.info(f"TERMS_FILTER_RESULT: Excluded sources breakdown: {excluded_sources}")
            logger.info(f"TERMS_FILTER_RESULT: Included sources breakdown: {included_sources}")
            logger.info(f"TERMS_FILTER_RESULT: After filtering: {len(filtered_by_terms)} documents remain (selected: {terms_profile})")
            
            filtered_results = filtered_by_terms
        else:
            logger.info(f"TERMS_FILTER_RESULT: No excluded patterns, keeping all {len(filtered_results)} documents")
    
    # Sort by score (descending), then prioritize selected terms documents, then by document name (ascending) for deterministic ordering
    def sort_key(result: Dict) -> tuple:
        score = result.get('score', 0)
        # Extract filename from source for consistent sorting
        source = result.get('source', '')
        # Extract just the filename if source is a path
        filename = os.path.basename(source) if source else ''
        filename_lower = filename.lower()
        
        # Check if this document matches the selected terms profile
        # Prioritize S3 bucket/path matching over filename to avoid false matches
        # (e.g., "reference standard contract terms and conditions.pdf" should NOT match)
        is_selected_terms = False
        if terms_profile:
            location = result.get('location', {})
            s3_location = location.get('s3Location', {})
            s3_uri = s3_location.get('uri', '').lower()
            s3_key = s3_location.get('key', '').lower()
            
            # Extract bucket name from S3 URI (format: s3://bucket-name/path/to/file.pdf)
            bucket_name_from_uri = ''
            if s3_uri and s3_uri.startswith('s3://'):
                try:
                    # Extract bucket name (everything between s3:// and first /)
                    bucket_name_from_uri = s3_uri.split('/')[2] if len(s3_uri.split('/')) > 2 else ''
                except:
                    bucket_name_from_uri = ''
            
            # Get allowed patterns for the selected terms profile
            allowed_patterns = terms_bucket_patterns.get(terms_profile, [])
            
            # Priority 1: Check S3 bucket name (most reliable indicator)
            # Bucket names like "onel-prod-general-terms" should match "general-terms" pattern
            if bucket_name_from_uri:
                bucket_match = any(pattern.lower() in bucket_name_from_uri for pattern in allowed_patterns)
                if bucket_match:
                    is_selected_terms = True
                    logger.debug(f"TERMS_MATCH: Matched via bucket name '{bucket_name_from_uri}' with pattern from {terms_profile}")
            
            # Priority 2: Check S3 key path (if bucket check didn't match)
            # Paths like "general-terms/document.pdf" should match
            if not is_selected_terms and s3_key:
                key_match = any(pattern.lower() in s3_key for pattern in allowed_patterns)
                if key_match:
                    is_selected_terms = True
                    logger.debug(f"TERMS_MATCH: Matched via S3 key path '{s3_key[:80]}' with pattern from {terms_profile}")
            
            # Priority 3: Check specific known filenames (only for exact matches)
            # This is for documents like "form_commonwealth-terms-and-conditions.pdf"
            if not is_selected_terms:
                # Only check specific filename patterns, not generic "terms" patterns
                specific_filename_patterns = [
                    'form_commonwealth-terms-and-conditions',  # General terms specific file
                    'updated it terms'  # IT terms updated specific file
                ]
                for pattern in allowed_patterns:
                    if pattern.lower() in specific_filename_patterns:
                        # Only match if the pattern appears as a complete word/phrase in the filename
                        if pattern.lower() in filename_lower:
                            is_selected_terms = True
                            logger.debug(f"TERMS_MATCH: Matched via specific filename pattern '{pattern}' in '{filename_lower}'")
                            break
            
            # Log if document was NOT matched (for debugging)
            if not is_selected_terms and terms_profile:
                logger.debug(f"TERMS_NO_MATCH: Document '{filename_lower[:80]}' from bucket '{bucket_name_from_uri}' did not match patterns for {terms_profile}")
        
        # Boost score for selected terms documents to ensure they appear prominently
        # Add a small boost (0.1) to selected terms documents so they rank higher
        # This ensures selected terms docs appear even if other docs have slightly higher scores
        boost_amount = 0.1 if is_selected_terms else 0.0
        boosted_score = score + boost_amount
        
        # Store boosted score in result for logging later
        result['_original_score'] = score
        result['_boosted_score'] = boosted_score
        result['_is_selected_terms'] = is_selected_terms
        
        # Return tuple: (negative boosted score for descending, filename for deterministic ordering)
        # The boost ensures selected terms docs come first unless other docs have significantly higher scores
        return (-boosted_score, filename_lower)
    
    sorted_results = sorted(filtered_results, key=sort_key)
    
    # Log score boosting and document ranking for debugging
    if terms_profile and sorted_results:
        # Count boosted documents
        boosted_count = sum(1 for r in sorted_results if r.get('_is_selected_terms', False))
        non_boosted_count = len(sorted_results) - boosted_count
        logger.info(f"SCORE_BOOST_SUMMARY: {boosted_count} documents boosted for selected profile '{terms_profile}' out of {len(sorted_results)} total documents")
        logger.info(f"SCORE_BOOST_SUMMARY: {non_boosted_count} documents NOT boosted (may be from other sources or don't match selected profile)")
        
        # Log each boosted document
        if boosted_count > 0:
            logger.info(f"SCORE_BOOST_DETAILS: Documents that received boost for '{terms_profile}':")
            for result in sorted_results:
                if result.get('_is_selected_terms', False):
                    original_score = result.get('_original_score', result.get('score', 0))
                    boosted_score = result.get('_boosted_score', original_score)
                    result_source = result.get('source', 'unknown')
                    result_location = result.get('location', {})
                    result_s3_uri = result_location.get('s3Location', {}).get('uri', '')
                    boost_amount = boosted_score - original_score
                    logger.info(f"SCORE_BOOST: '{result_source[:100]}' boosted from {original_score:.4f} to {boosted_score:.4f} (+{boost_amount:.4f}) (S3: {result_s3_uri[:80] if result_s3_uri else 'N/A'})")
        
        # Log top documents with their scores and ranking
        top_n = min(15, len(sorted_results))
        logger.info(f"DOCUMENT_RANKING: Top {top_n} documents after sorting (selected profile: {terms_profile}):")
        for idx, result in enumerate(sorted_results[:top_n], 1):
            result_source = result.get('source', 'unknown')
            original_score = result.get('_original_score', result.get('score', 0))
            boosted_score = result.get('_boosted_score', original_score)
            is_selected = result.get('_is_selected_terms', False)
            boost_indicator = "[BOOSTED]" if is_selected else "[NOT BOOSTED]"
            result_location = result.get('location', {})
            result_s3_uri = result_location.get('s3Location', {}).get('uri', '')
            
            # Extract bucket name from S3 URI for clarity
            bucket_name = 'unknown'
            if result_s3_uri:
                try:
                    # Extract bucket from s3://bucket-name/path
                    bucket_name = result_s3_uri.split('/')[2] if '/' in result_s3_uri else result_s3_uri
                except:
                    bucket_name = 'parse_error'
            
            boost_indicator = "⭐ BOOSTED" if is_selected else ""
            logger.info(f"  {idx}. Score: {original_score:.4f} → {boosted_score:.4f} {boost_indicator} | Source: {result_source[:50]} | Bucket: {bucket_name[:40]}")
        
        # Log high-scoring non-selected documents to understand why they rank high
        high_scoring_others = [r for r in sorted_results[:top_n] 
                              if not r.get('_is_selected_terms', False) and r.get('score', 0) > 0.7]
        if high_scoring_others:
            logger.info(f"HIGH_SCORE_ANALYSIS: {len(high_scoring_others)} non-selected documents with scores > 0.7:")
            for result in high_scoring_others[:5]:  # Log top 5 high-scoring others
                result_source = result.get('source', 'unknown')
                result_score = result.get('score', 0)
                result_location = result.get('location', {})
                result_s3_uri = result_location.get('s3Location', {}).get('uri', '')
                bucket_name = result_s3_uri.split('/')[2] if result_s3_uri and '/' in result_s3_uri else 'unknown'
                logger.info(f"  - Score: {result_score:.4f} | Source: {result_source[:50]} | Bucket: {bucket_name[:40]}")
    
    # Limit results for optimal performance
    return sorted_results[:min(max_results, OPTIMAL_RESULTS_PER_QUERY)]

def get_tool_definitions() -> List[Dict[str, Any]]:
    """Get tool definitions for Claude in Converse API format."""
    return [
        {
            "toolSpec": {
                "name": "retrieve_from_knowledge_base",
                "description": """
Exhaustively retrieve ALL relevant reference documents for conflict detection.

**When to use:**
- After analyzing vendor document structure
- When you need to check vendor language against Massachusetts requirements
- For each distinct section or topic area in vendor document

**Query Construction Best Practices:**
- Include specific contract terms from vendor document
- Add legal phrases and Massachusetts-specific terminology
- Include synonyms and variations of key terms
- Target 50-100+ unique terms per query
- Make queries distinct and non-overlapping

**Examples of Effective Queries:**
- "liability indemnity insurance Massachusetts ITS Terms and Conditions unlimited coverage"
- "payment terms invoicing net 30 days Massachusetts Commonwealth requirements"
- "data ownership intellectual property confidentiality Massachusetts state contracts"

**Output:**
Returns array of relevant document chunks with text, metadata, and relevance scores.
Optimized for maximum coverage with deduplication and relevance filtering.

Use 6-12+ targeted queries to ensure no conflicts are missed.
""",
                "inputSchema": {
                    "json": {
                        "type": "object",
                        "properties": {
                            "query": {
                                "type": "string",
                                "description": "Targeted search query to find reference documents. Use specific contract terms, legal phrases, or vendor-specific language. Try variations of important terms to catch all relevant content. Target 50-100+ unique terms per query."
                            },
                            "max_results": {
                                "type": "integer",
                                "description": "Maximum number of results to retrieve (auto-optimized for comprehensive coverage, default: 50)",
                                "default": 50
                            }
                        },
                        "required": ["query"]
                    }
                }
            }
        }
    ]

def _remove_uuid_prefix(filename: str) -> str:
    """
    Remove UUID prefix from filename if present.
    Format: uuid_filename.docx -> filename.docx
    """
    import re
    # Match UUID pattern (hex digits with optional hyphens) followed by underscore
    # Pattern: [a-f0-9-]+_filename.ext
    match = re.match(r'^[a-f0-9-]+_(.+)$', filename, re.IGNORECASE)
    return match.group(1) if match else filename

def _extract_source_from_result(metadata: Dict[str, Any], location: Dict[str, Any]) -> str:
    """
    Extract source document name from retrieval result metadata and location.
    
    Checks multiple possible fields where AWS Bedrock Knowledge Base might store
    the source information, in order of preference:
    1. metadata.source - primary source field
    2. location.s3Location.uri - full S3 URI
    3. location.s3Location.key - S3 key
    4. metadata.s3_location - alternative metadata field
    5. metadata.uri - URI in metadata
    6. Extract filename from S3 key if available
    
    Also removes UUID prefixes from filenames (format: uuid_filename.docx -> filename.docx)
    
    Args:
        metadata: Metadata dictionary from retrieval result
        location: Location dictionary from retrieval result
        
    Returns:
        Source document name or descriptive fallback if not found
    """
    import os
    
    # Try metadata.source first (primary field)
    if metadata.get('source'):
        return _remove_uuid_prefix(metadata['source'])
    
    # Try location.s3Location.uri (full S3 URI)
    s3_location = location.get('s3Location', {})
    if s3_location.get('uri'):
        uri = s3_location['uri']
        # Extract filename from URI (e.g., s3://bucket/path/file.pdf -> file.pdf)
        if '/' in uri:
            filename = uri.split('/')[-1]
            if filename:
                return _remove_uuid_prefix(filename)
    
    # Try location.s3Location.key (S3 key)
    if s3_location.get('key'):
        s3_key = s3_location['key']
        # Extract filename from S3 key (e.g., path/to/file.pdf -> file.pdf)
        if '/' in s3_key:
            filename = s3_key.split('/')[-1]
            if filename:
                return _remove_uuid_prefix(filename)
        return _remove_uuid_prefix(s3_key)
    
    # Try metadata.s3_location (alternative metadata field)
    if metadata.get('s3_location'):
        s3_loc = metadata['s3_location']
        if '/' in s3_loc:
            filename = s3_loc.split('/')[-1]
            if filename:
                return _remove_uuid_prefix(filename)
        return _remove_uuid_prefix(s3_loc)
    
    # Try metadata.uri
    if metadata.get('uri'):
        uri = metadata['uri']
        if '/' in uri:
            filename = uri.split('/')[-1]
            if filename:
                return _remove_uuid_prefix(filename)
        return _remove_uuid_prefix(uri)
    
    # Try extracting from any metadata field that looks like a path or filename
    for key, value in metadata.items():
        if isinstance(value, str) and ('/' in value or '.' in value):
            # Check if it looks like a file path
            if any(ext in value.lower() for ext in ['.pdf', '.docx', '.doc', '.txt', '.html']):
                filename = value.split('/')[-1] if '/' in value else value
                if filename and len(filename) > 3:  # Reasonable filename length
                    return _remove_uuid_prefix(filename)
    
    # If we have location info but no specific source, return a descriptive message
    if s3_location:
        bucket = s3_location.get('bucket', 'unknown-bucket')
        key = s3_location.get('key', 'unknown-key')
        # Extract filename from key if possible
        if '/' in key:
            filename = key.split('/')[-1]
            if filename:
                return _remove_uuid_prefix(filename)
        return _remove_uuid_prefix(key) if key != 'unknown-key' else f"{bucket}/{key}"
    
    # Last resort: log warning and return descriptive fallback
    logger.warning(f"Could not extract source from metadata: {metadata}, location: {location}")
    return 'Unknown Source'

def retrieve_from_knowledge_base(
    query: str, 
    max_results: int = 50,
    knowledge_base_id: str = None,
    region: str = None,
    terms_profile: str = None
) -> Dict[str, Any]:
    """
    Intelligently retrieve relevant documents from the knowledge base with optimization.
    
    Features:
    - Exponential backoff for throttling resilience
    - Semantic deduplication to prevent duplicate content
    - Intelligent chunking for optimal context usage
    - Relevance filtering to improve quality
    - Performance optimization for multiple queries
    
    Args:
        query: Search query for relevant documents
        max_results: Maximum number of results to return
        knowledge_base_id: Knowledge base ID from environment
        region: AWS region
        
    Returns:
        Dictionary containing optimized retrieved documents and metadata
    """
    
    def _retrieve_with_retry(retry_count: int = 0) -> Dict[str, Any]:
        """Internal function to handle retrieval with exponential backoff."""
        try:

            
            response = bedrock_agent_client.retrieve(
                knowledgeBaseId=knowledge_base_id,
                retrievalQuery={'text': query},
                retrievalConfiguration={
                    'vectorSearchConfiguration': {
                        'numberOfResults': max_results
                    }
                }
            )
            
            return {"success": True, "response": response, "retry_count": retry_count}
            
        except Exception as e:
            error_msg = str(e)
            
            # Check for throttling errors and implement exponential backoff
            if ("ThrottlingException" in error_msg or "Too many tokens" in error_msg or 
                "rate" in error_msg.lower() or "throttl" in error_msg.lower()):
                
                if retry_count < MAX_RETRIES:
                    delay = min(BASE_DELAY * (BACKOFF_MULTIPLIER ** retry_count), MAX_DELAY)

                    time.sleep(delay)
                    return _retrieve_with_retry(retry_count + 1)
                else:

                    return {
                        "success": False,
                        "error": f"Throttling limit exceeded after {MAX_RETRIES} retries",
                        "retry_count": retry_count
                    }
            else:
                # Non-throttling error
                return {"success": False, "error": error_msg, "retry_count": retry_count}
    
    try:
        # Get knowledge base ID from environment if not provided
        if not knowledge_base_id:
            knowledge_base_id = os.environ.get('KNOWLEDGE_BASE_ID')
        
        if not knowledge_base_id:
            return {
                "success": False,
                "error": "Knowledge base ID not available",
                "results": []
            }
        
        # Check query cache to avoid duplicate requests in same session
        # Include knowledge_base_id in cache key to prevent cross-KB cache pollution
        cache_key = f"{knowledge_base_id}:{_calculate_content_signature(query)}"
        if cache_key in _query_cache:

            cached_result = _query_cache[cache_key].copy()
            cached_result["cached"] = True
            logger.info(f"QUERY_CACHE_HIT: Using cached results for KB={knowledge_base_id}, query_hash={cache_key.split(':')[1]}")
            return cached_result
        
        # Execute retrieval with retry logic
        retrieval_result = _retrieve_with_retry()
        
        if not retrieval_result["success"]:
            error_response = {
                "success": False,
                "error": retrieval_result["error"],
                "query": query,
                "results": [],
                "retry_count": retrieval_result.get("retry_count", 0)
            }
            return error_response
        
        response = retrieval_result["response"]
        retry_count = retrieval_result["retry_count"]
        
        # Process and optimize the results
        raw_results = []
        source_documents = set()
        for result in response.get('retrievalResults', []):
            content = result.get('content', {})
            metadata = result.get('metadata', {})
            location = result.get('location', {})
            
            # Extract source from multiple possible locations
            source = _extract_source_from_result(metadata, location)
            source_documents.add(source)
            
            raw_results.append({
                "text": content.get('text', ''),
                "score": result.get('score', 0),
                "source": source,
                "metadata": metadata,
                "location": location  # Include location to check S3 URI for bucket name
            })
        
        # Enhanced logging: Track which reference documents were found
        logger.info(f"KNOWLEDGE_BASE_QUERY: '{query[:100]}...' found {len(raw_results)} results from {len(source_documents)} source documents: {list(source_documents)}")
        
        # Apply intelligent filtering and prioritization
        filtered_results = _filter_and_prioritize_results(raw_results, max_results, terms_profile=terms_profile)
        
        # Process results with chunking and deduplication
        optimized_results = []
        duplicates_filtered = 0
        chunks_created = 0
        
        for result in filtered_results:
            # Check for duplicate content
            if _is_duplicate_content(result["text"]):
                duplicates_filtered += 1
                continue
            
            # Apply intelligent chunking for large content
            chunks = _chunk_content_intelligently(result["text"])
            
            if len(chunks) == 1:
                # Single chunk, add as-is with optimization metadata
                optimized_result = result.copy()
                optimized_result.update({
                    "tokens_estimated": len(result["text"]) // 4,
                    "optimized": True,
                    "chunk_info": {"is_chunked": False, "total_chunks": 1}
                })
                optimized_results.append(optimized_result)
            else:
                # Multiple chunks, add each with metadata
                chunks_created += len(chunks) - 1  # -1 because we count additional chunks
                for i, chunk in enumerate(chunks):
                    chunk_result = result.copy()
                    chunk_result.update({
                        "text": chunk,
                        "tokens_estimated": len(chunk) // 4,
                        "optimized": True,
                        "chunk_info": {
                            "is_chunked": True,
                            "chunk_number": i + 1,
                            "total_chunks": len(chunks),
                            "original_source": result["source"]
                        }
                    })
                    optimized_results.append(chunk_result)
        
        # Generate comprehensive response with optimization statistics
        final_response = {
            "success": True,
            "query": query,
            "results_count": len(optimized_results),
            "results": optimized_results,
            "optimization_stats": {
                "raw_results_retrieved": len(raw_results),
                "filtered_by_relevance": len(raw_results) - len(filtered_results),
                "duplicates_filtered": duplicates_filtered,
                "chunks_created": chunks_created,
                "final_optimized_count": len(optimized_results),
                "retry_count": retry_count,
                "avg_relevance_score": sum(r.get('score', 0) for r in optimized_results) / len(optimized_results) if optimized_results else 0,
                "cache_hit": False
            },
            "performance_metrics": {
                "query_hash": cache_key.split(':')[1],  # Extract just the query hash part
                "cache_key": cache_key,
                "processing_successful": True,
                "optimization_ratio": f"{len(optimized_results)}/{len(raw_results)}" if raw_results else "0/0"
            }
        }
        
        # Cache the result for future use in this session (keyed by KB ID + query)
        _query_cache[cache_key] = final_response.copy()
        logger.info(f"QUERY_CACHE_STORE: Cached results for KB={knowledge_base_id}, query_hash={cache_key.split(':')[1]}")
        


        
        return final_response
        
    except Exception as e:
        logger.error(f"Error in optimized knowledge base retrieval: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "query": query,
            "results": [],
            "optimization_stats": {
                "error_occurred": True,
                "error_type": type(e).__name__
            }
        }

def clear_knowledge_base_cache():
    """
    Clear the session cache for knowledge base retrieval.
    Call this between document review sessions to ensure fresh retrievals.
    """
    global _content_cache, _query_cache
    _content_cache.clear()
    _query_cache.clear()


def get_cache_statistics() -> Dict[str, Any]:
    """
    Get current cache statistics for monitoring and debugging.
    
    Returns:
        Dictionary containing cache usage statistics
    """
    return {
        "content_cache_size": len(_content_cache),
        "query_cache_size": len(_query_cache),
        "cache_memory_usage": {
            "content_signatures": len(_content_cache),
            "cached_queries": len(_query_cache)
        }
    }


def redline_document(
    analysis_data: str,
    document_s3_key: str,
    bucket_type: str = "user_documents",
    session_id: str = None,
    user_id: str = None
) -> Dict[str, Any]:
    """
    Complete redlining workflow: download document, extract content, apply redlining, upload result.
    Handles DOCX documents with text modification redlining.
    
    Args:
        analysis_data: Analysis text containing conflict information
        document_s3_key: S3 key of the original document
        bucket_type: Type of source bucket (user_documents, knowledge, agent_processing)
        session_id: Session ID for organizing output files
        user_id: User ID for organizing output files
        
    Returns:
        Dictionary containing redlined document information and processing results
    """
    
    try:
        logger.info(f"REDLINE_START: Document={document_s3_key}, Analysis={len(analysis_data)} chars")
        
        logger.info("FILE_TYPE_DETECTED: DOCX")
        
        # Get bucket configurations
        source_bucket = _get_bucket_name(bucket_type)
        agent_processing_bucket = os.environ.get('AGENT_PROCESSING_BUCKET')
        
        if not agent_processing_bucket or not source_bucket:
            # Cleanup on failure
            if session_id and user_id:
                try:
                    _cleanup_session_documents(session_id, user_id)
                except Exception as cleanup_error:
                    logger.error(f"Session cleanup error during bucket config check: {cleanup_error}")
            return {
                "success": False,
                "error": "Required buckets not configured"
            }
        

        
        # Step 1: Copy document to agent processing bucket
        agent_document_key = _copy_document_to_processing(document_s3_key, source_bucket, agent_processing_bucket)
        
        # Step 3: Parse conflicts and create redline items from analysis data
        redline_items = parse_conflicts_for_redlining(analysis_data)
        logger.info(f"REDLINE_PARSE: Found {len(redline_items)} conflicts to redline")
        if redline_items:
            logger.info(f"REDLINE_PARSE: First conflict preview: '{redline_items[0].get('text', '')[:100]}...'")
        
        # Step 3.5: Assign sequential serial numbers based on comment content deduplication
        # Track seen comments by normalized content to prevent duplicates and assign serial numbers
        seen_comments = {}  # Key: normalized comment content, Value: serial_number
        serial_number = 0
        
        for item in redline_items:
            comment = item.get('comment', '').strip()
            if not comment:
                continue
            
            # Normalize comment content for duplicate detection (remove serial numbers if present)
            # Extract the core content: everything after "CONFLICT" but before any existing serial number
            # Preserve the Reference section for normalization
            normalized_comment = comment
            # Remove any existing conflict ID pattern to get core content (but keep Reference section)
            normalized_comment = re.sub(r'^CONFLICT\s+\S+\s*\(', 'CONFLICT (', normalized_comment)
            normalized_comment = normalized_comment.strip()
            
            # Check if we've seen this comment content before
            if normalized_comment not in seen_comments:
                serial_number += 1
                seen_comments[normalized_comment] = serial_number
                # Update comment with serial number
                # Format: "CONFLICT {serial_number} ({type}): {rationale}\n\nReference: {source_doc} ({clause_ref})"
                # Extract type and everything after (including rationale and Reference section)
                match = re.match(r'^CONFLICT\s+\S+\s*\(([^)]+)\):\s*(.+)', comment, re.DOTALL)
                if match:
                    conflict_type = match.group(1)
                    rationale_and_ref = match.group(2)  # This includes rationale + Reference section
                    item['comment'] = f"CONFLICT {serial_number} ({conflict_type}): {rationale_and_ref}"
                else:
                    # Fallback: just prepend serial number, preserve entire comment including Reference
                    item['comment'] = f"CONFLICT {serial_number}: {comment.replace('CONFLICT ', '').replace('CONFLICT', '')}"
                item['serial_number'] = serial_number
                logger.debug(f"Assigned serial number {serial_number} to conflict with comment: {normalized_comment[:100]}...")
            else:
                # Duplicate comment - use existing serial number
                existing_serial = seen_comments[normalized_comment]
                match = re.match(r'^CONFLICT\s+\S+\s*\(([^)]+)\):\s*(.+)', comment, re.DOTALL)
                if match:
                    conflict_type = match.group(1)
                    rationale_and_ref = match.group(2)  # This includes rationale + Reference section
                    item['comment'] = f"CONFLICT {existing_serial} ({conflict_type}): {rationale_and_ref}"
                else:
                    item['comment'] = f"CONFLICT {existing_serial}: {comment.replace('CONFLICT ', '').replace('CONFLICT', '')}"
                item['serial_number'] = existing_serial
                logger.info(f"DUPLICATE_COMMENT: Found duplicate comment content, reusing serial number {existing_serial}")
        
        logger.info(f"REDLINE_SERIAL: Assigned {serial_number} unique serial numbers to {len(redline_items)} conflicts (based on comment content deduplication)")
        
        if not redline_items:
            # Cleanup reference documents when no conflicts detected
            cleanup_result = None
            if session_id and user_id:
                try:
                    cleanup_result = _cleanup_session_documents(session_id, user_id)
                    logger.info(f"Cleanup completed for 0 conflicts case: {cleanup_result}")
                except Exception as cleanup_error:
                    logger.error(f"Session cleanup error during no conflicts check: {cleanup_error}")
            return {
                "success": True,
                "no_conflicts": True,
                "message": "No conflicts found in analysis data for redlining",
                "cleanup_performed": cleanup_result is not None,
                "cleanup_result": cleanup_result
            }
        
        # DOCX Processing
        logger.info("PROCESSING_DOCX: Using DOCX text modification redlining")
        
        # DOCX Processing - Original code path
        # Step 2: Download and load the DOCX document
        doc = _download_and_load_document(agent_processing_bucket, agent_document_key)
        
        # Debug logging: Log document structure for troubleshooting
        logger.info(f"DOCUMENT_DEBUG: Loaded document with {len(doc.paragraphs)} paragraphs")
        for i, para in enumerate(doc.paragraphs[:5]):  # Log first 5 paragraphs
            if para.text.strip():
                logger.info(f"DOCUMENT_DEBUG: Para {i}: '{para.text[:100]}...'")
        
        # Note: redline_items was already parsed above, no need to parse again
        # If we reach here, redline_items should not be empty (checked earlier)
        # But add a safety check to prevent crashes
        if not redline_items:
            # This should not happen since we checked earlier, but handle gracefully
            logger.warning("REDLINE_WARNING: redline_items is empty after document load, this should not happen")
            cleanup_result = None
            if session_id and user_id:
                try:
                    cleanup_result = _cleanup_session_documents(session_id, user_id)
                    logger.info(f"Cleanup completed for unexpected 0 conflicts case: {cleanup_result}")
                except Exception as cleanup_error:
                    logger.error(f"Session cleanup error during unexpected no conflicts check: {cleanup_error}")
            return {
                "success": True,
                "no_conflicts": True,
                "message": "No conflicts found in analysis data for redlining",
                "cleanup_performed": cleanup_result is not None,
                "cleanup_result": cleanup_result
            }
        
        # Log first conflict preview safely
        if redline_items:
            logger.info(f"REDLINE_PARSE: First conflict preview: '{redline_items[0].get('text', '')[:100]}...'")
        
        # Step 4: Apply redlining with exact sentence matching
        logger.info(f"REDLINE_APPLY: Starting redlining - {len(redline_items)} conflicts, {len(doc.paragraphs)} paragraphs")
        
        results = apply_exact_sentence_redlining(doc, redline_items)
        logger.info(f"REDLINE_RESULTS: Matches found: {results['matches_found']}")
        logger.info(f"REDLINE_RESULTS: Failed matches: {len(results.get('failed_matches', []))}")
        logger.info(f"REDLINE_RESULTS: Success rate: {(results['matches_found']/len(redline_items)*100):.1f}%")

        # Step 5: Save and upload redlined document
        redlined_s3_key = _create_redlined_filename(agent_document_key, session_id, user_id)
        upload_success = _save_and_upload_document(doc, agent_processing_bucket, redlined_s3_key, {
            'original_document': document_s3_key,
            'agent_document': agent_document_key,
            'redlined_by': 'Legal-AI',
            'conflicts_count': str(len(redline_items)),
            'matches_found': str(results['matches_found'])
        })
        logger.info(f"REDLINE_UPLOAD: Uploaded redlined document to {redlined_s3_key}")
        
        if not upload_success:
            # Cleanup on failure
            if session_id and user_id:
                try:
                    _cleanup_session_documents(session_id, user_id)
                except Exception as cleanup_error:
                    logger.error(f"Session cleanup error during upload failure: {cleanup_error}")
            return {
                "success": False,
                "error": "Failed to upload redlined document"
            }
        
        # Original redlining completion
        result = {
            "success": True,
            "original_document": document_s3_key,
            "agent_document": agent_document_key,
            "redlined_document": redlined_s3_key,
            "conflicts_processed": len(redline_items),
            "matches_found": results['matches_found'],
            "paragraphs_with_redlines": results['paragraphs_with_redlines'],
            "bucket": agent_processing_bucket,
            "message": f"Successfully created redlined document with {results['matches_found']} conflicts highlighted"
        }

        # NEW: After successful redlining, cleanup session documents
        if session_id and user_id:

            try:
                cleanup_result = _cleanup_session_documents(session_id, user_id)
                result["cleanup_performed"] = True
                result["cleanup_result"] = cleanup_result
                
                if cleanup_result.get('success'):
                    pass
                else:
                    pass
                    
            except Exception as cleanup_error:
                # Don't fail redlining if cleanup fails
                logger.error(f"Session cleanup error: {cleanup_error}")
                result["cleanup_performed"] = False
                result["cleanup_error"] = str(cleanup_error)
        else:

            result["cleanup_performed"] = False

        return result
        
    except Exception as e:
        logger.error(f"Error in redlining workflow: {str(e)}")
        # Cleanup on failure
        if session_id and user_id:
            try:
                _cleanup_session_documents(session_id, user_id)
            except Exception as cleanup_error:
                logger.error(f"Session cleanup error during redlining workflow exception: {cleanup_error}")
        return {
            "success": False,
            "error": str(e),
            "original_document": document_s3_key
        }


def parse_conflicts_for_redlining(analysis_data: str) -> List[Dict[str, str]]:
    """
    Parse conflicts data for redlining. Supports both JSON object format (with explanation and conflicts),
    JSON array format, and markdown table format.
    Extracts exact sentences from vendor_quote field that should be present in vendor document.
    
    Args:
        analysis_data: Analysis string containing JSON object/array or markdown table
        
    Returns:
        List of redline items with exact text to match and comments
    """
    logger.info(f"PARSE_START: Analysis data length: {len(analysis_data)} characters")
    logger.info(f"PARSE_START: Analysis preview: {analysis_data[:150]}...")

    redline_items = []
    explanation = ""
    
    try:
        # Try to parse as JSON first (new format)
        import json
        import re
        
        # First, try to find JSON object pattern (new format with explanation and conflicts)
        json_object_match = re.search(r'\{[\s\S]*"conflicts"[\s\S]*\}', analysis_data)
        if json_object_match:
            try:
                json_str = json_object_match.group(0)
                parsed_json = json.loads(json_str)
                
                # Check if it's the new structure with explanation and conflicts
                if isinstance(parsed_json, dict) and "conflicts" in parsed_json:
                    explanation = parsed_json.get("explanation", "")
                    conflicts_json = parsed_json.get("conflicts", [])
                    
                    # Log the explanation
                    if explanation:
                        logger.info(f"PARSE_EXPLANATION: {explanation}")
                    else:
                        logger.info("PARSE_EXPLANATION: No explanation provided")
                    
                    if isinstance(conflicts_json, list):
                        logger.info(f"PARSE_JSON: Found JSON object with explanation and {len(conflicts_json)} conflicts")
                    else:
                        logger.warning(f"PARSE_JSON: Expected conflicts to be a list, got {type(conflicts_json)}")
                        conflicts_json = []
                else:
                    # Not the expected structure, fall through to array parsing
                    conflicts_json = None
            except json.JSONDecodeError as e:
                logger.warning(f"PARSE_JSON_OBJECT_FAILED: Could not parse as JSON object, trying array format: {e}")
                conflicts_json = None
        else:
            conflicts_json = None
        
        # If we didn't find the new structure, try to find JSON array pattern (backwards compatibility)
        if conflicts_json is None:
            json_match = re.search(r'\[[\s\S]*\]', analysis_data)
            if json_match:
                try:
                    json_str = json_match.group(0)
                    conflicts_json = json.loads(json_str)
                    
                    if isinstance(conflicts_json, list):
                        logger.info(f"PARSE_JSON: Found JSON array with {len(conflicts_json)} conflicts (backwards compatibility mode)")
                    else:
                        conflicts_json = None
                except json.JSONDecodeError as e:
                    logger.warning(f"PARSE_JSON_ARRAY_FAILED: Could not parse as JSON array: {e}")
                    conflicts_json = None
        
        # Process conflicts if we found them
        if conflicts_json is not None and isinstance(conflicts_json, list):
                    
                    validated_count = 0
                    validation_errors = []
                    
                    for idx, conflict in enumerate(conflicts_json):
                        if not isinstance(conflict, dict):
                            logger.warning(f"PARSE_SKIP: Conflict {idx} is not a dict, skipping")
                            continue
                        
                        # Log raw conflict data before processing
                        conflict_id = conflict.get('clarification_id', f'Unknown-{idx}')
                        raw_vendor_quote = conflict.get('vendor_quote', '')
                        logger.info(f"PARSE_RAW_CONFLICT: ID={conflict_id}, vendor_quote type={type(raw_vendor_quote)}, length={len(str(raw_vendor_quote))}")
                        logger.info(f"PARSE_RAW_CONFLICT: vendor_quote (first 200 chars)='{str(raw_vendor_quote)[:200]}...'")
                        if '\\' in str(raw_vendor_quote):
                            logger.info(f"PARSE_RAW_CONFLICT: vendor_quote contains backslashes, count={str(raw_vendor_quote).count('\\\\')}")
                        
                        # Use Pydantic validation if available
                        if PYDANTIC_AVAILABLE and ConflictModel:
                            try:
                                validated_conflict = ConflictModel(**conflict)
                                
                                # Log after Pydantic validation
                                logger.info(f"PARSE_AFTER_PYDANTIC: ID={validated_conflict.clarification_id}, vendor_quote type={type(validated_conflict.vendor_quote)}, length={len(validated_conflict.vendor_quote)}")
                                logger.info(f"PARSE_AFTER_PYDANTIC: vendor_quote (first 200 chars)='{validated_conflict.vendor_quote[:200]}...'")
                                
                                # Use comment format: CONFLICT ID (type): rationale
                                comment = f"CONFLICT {validated_conflict.clarification_id} ({validated_conflict.conflict_type}): {validated_conflict.rationale}"
                                
                                # Add reference section
                                if validated_conflict.source_doc and validated_conflict.source_doc.strip():
                                    comment += f"\n\nReference: {validated_conflict.source_doc.strip()}"
                                    if validated_conflict.clause_ref and validated_conflict.clause_ref.strip() and validated_conflict.clause_ref.lower() not in ['n/a', 'na', 'none']:
                                        comment += f" ({validated_conflict.clause_ref.strip()})"
                                
                                # Store original vendor_quote (only normalize escaped quotes for JSON parsing)
                                # Full normalization will happen during matching to preserve original structure
                                vendor_quote_text = normalize_escaped_quotes(validated_conflict.vendor_quote.strip())
                                
                                # Validate for truncated quotes
                                is_truncated = _validate_vendor_quote_completeness(vendor_quote_text, validated_conflict.clarification_id)
                                if is_truncated:
                                    logger.warning(f"PARSE_VALIDATION: ID={validated_conflict.clarification_id}, vendor_quote appears to be TRUNCATED/INCOMPLETE. Length={len(vendor_quote_text)}, ends_with='{vendor_quote_text[-30:]}'")
                                
                                redline_items.append({
                                    'text': vendor_quote_text,
                                    'comment': comment,
                                    'author': 'Legal-AI',
                                    'initials': 'LAI',
                                    'clarification_id': validated_conflict.clarification_id,
                                    'conflict_type': validated_conflict.conflict_type,
                                    'source_doc': validated_conflict.source_doc,
                                    'clause_ref': validated_conflict.clause_ref,
                                    'summary': validated_conflict.summary,
                                    'rationale': validated_conflict.rationale
                                })
                                validated_count += 1
                                
                            except Exception as e:
                                error_msg = f"Conflict {idx} validation failed: {str(e)}"
                                validation_errors.append(error_msg)
                                logger.warning(f"PARSE_VALIDATION_ERROR: {error_msg}")
                                # Fall through to manual validation for backwards compatibility
                                pass
                        
                        # Fallback to manual validation if Pydantic not available or validation failed
                        if not (PYDANTIC_AVAILABLE and ConflictModel):
                            clarification_id = conflict.get('clarification_id', '')
                            vendor_quote = conflict.get('vendor_quote', '')
                            summary = conflict.get('summary', '')
                            source_doc = conflict.get('source_doc', '')
                            clause_ref = conflict.get('clause_ref', '')
                            conflict_type = conflict.get('conflict_type', '')
                            rationale = conflict.get('rationale', '')
                            
                            logger.info(f"PARSE_MANUAL: ID={clarification_id}, vendor_quote type={type(vendor_quote)}, length={len(str(vendor_quote))}")
                            logger.info(f"PARSE_MANUAL: vendor_quote (first 200 chars)='{str(vendor_quote)[:200]}...'")
                            
                            # Clean up vendor quote - remove surrounding quotes if present
                            vendor_quote_clean = str(vendor_quote).strip()
                            if vendor_quote_clean.startswith('"') and vendor_quote_clean.endswith('"'):
                                logger.info(f"PARSE_MANUAL: ID={clarification_id}, removing surrounding quotes")
                                vendor_quote_clean = vendor_quote_clean[1:-1]
                            
                            # Only normalize escaped quotes - full normalization happens during matching
                            vendor_quote_before_norm = vendor_quote_clean
                            vendor_quote_clean = normalize_escaped_quotes(vendor_quote_clean)
                            logger.info(f"PARSE_MANUAL_NORM: ID={clarification_id}, before='{vendor_quote_before_norm[:150]}...', after='{vendor_quote_clean[:150]}...'")
                            if vendor_quote_before_norm != vendor_quote_clean:
                                logger.info(f"PARSE_MANUAL_NORM: ID={clarification_id}, normalization changed the text")
                            
                            # Create redline item using exact vendor quote for matching
                            if vendor_quote_clean.strip():  # Only add if we have actual text
                                # Use comment format: CONFLICT ID (type): rationale
                                comment = f"CONFLICT {clarification_id} ({conflict_type}): {rationale}"
                                
                                # Add reference section if source_doc is provided
                                if source_doc and str(source_doc).strip():
                                    comment += f"\n\nReference: {str(source_doc).strip()}"
                                    if clause_ref and str(clause_ref).strip() and str(clause_ref).lower() not in ['n/a', 'na', 'none']:
                                        comment += f" ({str(clause_ref).strip()})"
                                
                                redline_items.append({
                                    'text': vendor_quote_clean.strip(),
                                    'comment': comment,
                                    'author': 'Legal-AI',
                                    'initials': 'LAI',
                                    'clarification_id': str(clarification_id),
                                    'conflict_type': str(conflict_type),
                                    'source_doc': str(source_doc),
                                    'clause_ref': str(clause_ref),
                                    'summary': str(summary),
                                    'rationale': str(rationale)
                                })
                    
                    if PYDANTIC_AVAILABLE and ConflictModel:
                        logger.info(f"PARSE_VALIDATION: Validated {validated_count}/{len(conflicts_json)} conflicts using Pydantic")
                        if validation_errors:
                            logger.warning(f"PARSE_VALIDATION: {len(validation_errors)} validation errors occurred")
                    
                    # Deduplicate and filter using composite key (clarification_id + vendor_quote text)
                    # This ensures conflicts with same ID but different text (from different chunks) are kept
                    seen_conflicts = {}  # Key: (clarification_id, normalized_text)
                    deduplicated_items = []
                    for item in redline_items:
                        clarification_id = item.get('clarification_id', '')
                        text_val = (item.get('text') or '').strip()
                        # Filter placeholders/empty like 'N/A' or too short strings
                        if not text_val or text_val.lower() in ['n/a', 'na', 'none', 'n.a.', 'n.a', 'not available'] or len(text_val) < 5:
                            logger.warning(f"PARSE_FILTER: Skipping placeholder/empty conflict for ID={clarification_id} text='{text_val}'")
                            continue
                        
                        # Create composite key: clarification_id + normalized text (first 100 chars for uniqueness)
                        normalized_text = text_val.lower().strip()[:100]  # Normalize and truncate for key
                        composite_key = (clarification_id, normalized_text)
                        
                        if composite_key not in seen_conflicts:
                            seen_conflicts[composite_key] = True
                            deduplicated_items.append(item)
                        else:
                            logger.warning(f"PARSE_DEDUP: Duplicate conflict found - ID={clarification_id}, text_preview='{text_val[:50]}...', skipping")
                    
                    logger.info(f"PARSE_COMPLETE: Parsed {len(redline_items)} conflicts from JSON, {len(deduplicated_items)} unique conflicts after deduplication (using composite key: clarification_id + text)")
                    if len(redline_items) != len(deduplicated_items):
                        dropped = len(redline_items) - len(deduplicated_items)
                        logger.info(f"PARSE_DEDUP_SUMMARY: Dropped {dropped} duplicate conflicts (same ID + text combination)")
                    
                    # Store explanation in a way that can be accessed later if needed
                    # For now, we just log it - conflicts are what we return
                    return deduplicated_items
        
        # Fallback to markdown table parsing (backwards compatibility)
        lines = analysis_data.strip().split('\n')
        
        # Find the table header and data rows
        header_found = False
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines and markdown formatting
            if not line or line.startswith('|---') or line.startswith('|-'):
                continue
            
            # Check if this is the header row (should contain all expected columns)
            if ('| Clarification ID |' in line and '| Vendor Quote |' in line and 
                '| Summary |' in line and '| Rationale |' in line):
                header_found = True
                continue
            
            # Process data rows
            if header_found and line.startswith('|') and line.endswith('|'):
                # Split by pipe and clean up
                parts = [part.strip() for part in line.split('|')[1:-1]]
                
                if len(parts) >= 7:  # Need 7 columns for the complete table
                    clarification_id = parts[0]
                    vendor_quote = parts[1]  # This contains the exact text from vendor document for redlining
                    summary = parts[2]  # Plain-language context
                    source_doc = parts[3]
                    clause_ref = parts[4]
                    conflict_type = parts[5]
                    rationale = parts[6]
                    
                    # Clean up vendor quote - remove surrounding quotes if present
                    vendor_quote_clean = vendor_quote.strip()
                    if vendor_quote_clean.startswith('"') and vendor_quote_clean.endswith('"'):
                        vendor_quote_clean = vendor_quote_clean[1:-1]
                    
                    # Only normalize escaped quotes - full normalization happens during matching
                    vendor_quote_clean = normalize_escaped_quotes(vendor_quote_clean)
                    
                    # Create redline item using exact vendor quote for matching
                    if vendor_quote_clean.strip():  # Only add if we have actual text
                        # Use Ritik's comment format from main branch: CONFLICT ID (type): rationale
                        comment = f"CONFLICT {clarification_id} ({conflict_type}): {rationale}"
                        
                        # Add reference section if source_doc is provided
                        if source_doc and source_doc.strip():
                            comment += f"\n\nReference: {source_doc.strip()}"
                            if clause_ref and clause_ref.strip() and clause_ref.lower() not in ['n/a', 'na', 'none']:
                                comment += f" ({clause_ref.strip()})"
                        
                        redline_items.append({
                            'text': vendor_quote_clean.strip(),  # Exact sentence from vendor document
                            'comment': comment,
                            'author': 'Legal-AI',
                            'initials': 'LAI',
                            'clarification_id': clarification_id,
                            'conflict_type': conflict_type,
                            'source_doc': source_doc,
                            'clause_ref': clause_ref,
                            'summary': summary,
                            'rationale': rationale
                        })
                        
        # Deduplicate conflicts using composite key (clarification_id + vendor_quote text)
        # This ensures conflicts with same ID but different text (from different chunks) are kept
        seen_conflicts = {}  # Key: (clarification_id, normalized_text)
        deduplicated_items = []
        for item in redline_items:
            clarification_id = item.get('clarification_id', '')
            text_val = (item.get('text') or '').strip()
            # Only filter truly empty or placeholder values - be more permissive with short text
            # User preference: better to include conflicts than lose them
            if not text_val or text_val.lower() in ['n/a', 'na', 'none', 'n.a.', 'n.a', 'not available']:
                logger.warning(f"PARSE_FILTER: Skipping placeholder conflict for ID={clarification_id} text='{text_val}'")
                continue
            # Allow very short text (minimum 3 chars instead of 5) - user wants to preserve conflicts
            if len(text_val) < 3:
                logger.warning(f"PARSE_FILTER: Skipping extremely short conflict for ID={clarification_id} text='{text_val}' (length: {len(text_val)})")
                continue
            
            # Create composite key: clarification_id + normalized text (first 100 chars for uniqueness)
            normalized_text = text_val.lower().strip()[:100]  # Normalize and truncate for key
            composite_key = (clarification_id, normalized_text)
            
            if composite_key not in seen_conflicts:
                seen_conflicts[composite_key] = True
                deduplicated_items.append(item)
            else:
                logger.warning(f"PARSE_DEDUP: Duplicate conflict found - ID={clarification_id}, text_preview='{text_val[:50]}...', skipping")
        
        logger.info(f"PARSE_COMPLETE: Parsed {len(redline_items)} conflicts from analysis, {len(deduplicated_items)} unique conflicts after deduplication (using composite key: clarification_id + text)")
        if len(redline_items) != len(deduplicated_items):
            dropped = len(redline_items) - len(deduplicated_items)
            logger.info(f"PARSE_DEDUP_SUMMARY: Dropped {dropped} duplicate conflicts (same ID + text combination)")
        for i, item in enumerate(deduplicated_items[:2]):
            logger.info(f"PARSE_CONFLICT_{i+1}: ID={item.get('clarification_id')}, Text='{item.get('text', '')[:60]}...'")

        # Return deduplicated list
        redline_items = deduplicated_items
        
    except Exception as e:
        logger.error(f"Error parsing conflicts for redlining: {str(e)}")
    
    return redline_items


def _validate_vendor_quote_completeness(vendor_quote: str, clarification_id: str) -> bool:
    """
    Validate if vendor quote appears to be complete or truncated.
    
    Checks for signs that the quote was cut off mid-sentence or is incomplete.
    
    Args:
        vendor_quote: The vendor quote text to validate
        clarification_id: ID of the conflict for logging
        
    Returns:
        True if quote appears truncated, False if complete
    """
    if not vendor_quote or len(vendor_quote.strip()) < 10:
        return True  # Too short to be meaningful
    
    stripped = vendor_quote.rstrip()
    
    # Check if quote ends with proper punctuation (complete sentence/clause)
    ends_with_punctuation = stripped.endswith(('.', '!', '?', ';', ':', ')', ']', '}'))
    
    # Check if quote ends with common connecting words (likely truncated)
    common_end_words = ('or', 'and', 'the', 'a', 'an', 'to', 'of', 'in', 'for', 'with', 'from', 'by', 'at', 'on', 'as', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can')
    ends_with_common_word = any(stripped.lower().endswith(f' {word}') or stripped.lower().endswith(f', {word}') or stripped.lower().endswith(f'; {word}') for word in common_end_words)
    
    # Check if quote is suspiciously short (less than 100 chars is often incomplete for legal clauses)
    is_short = len(vendor_quote) < 100
    
    # Check if quote ends mid-word or with incomplete phrase
    ends_mid_phrase = (
        not ends_with_punctuation and 
        (ends_with_common_word or is_short)
    )
    
    is_truncated = ends_mid_phrase or (is_short and not ends_with_punctuation)
    
    if is_truncated:
        logger.warning(f"VALIDATE_TRUNCATED: ID={clarification_id}, quote appears truncated. "
                      f"Length={len(vendor_quote)}, ends_with_punctuation={ends_with_punctuation}, "
                      f"ends_with_common_word={ends_with_common_word}, is_short={is_short}, "
                      f"last_30_chars='{vendor_quote[-30:]}'")
    
    return is_truncated


def normalize_escaped_quotes(text: str) -> str:
    """
    Normalize escaped quotes in vendor_quote to match document text.
    
    Handles cases where vendor_quote contains literal \" or \' that need to match
    document text with regular " or ' quotes.
    
    Since json.loads() already handles JSON escapes, if we see \"
    it means literal backslash+quote. We normalize to just quote.
    
    Args:
        text: Text that may contain escaped quotes
        
    Returns:
        Text with escaped quotes normalized to regular quotes
    """
    if not text:
        return text
    
    # Replace literal \" with " (backslash followed by quote)
    # This handles cases like: \"word\" -> "word"
    text = text.replace('\\"', '"')
    # Replace literal \' with ' (backslash followed by single quote)
    text = text.replace("\\'", "'")
    # Handle unicode escapes (though these should be handled by JSON parsing)
    text = text.replace('\\u0022', '"')  # Unicode double quote
    text = text.replace('\\u0027', "'")  # Unicode single quote
    
    return text


def normalize_quotes(text: str) -> str:
    """
    Normalize various quote characters to standard ASCII quotes.
    
    Handles curly quotes, smart quotes, and other quote variants
    that may differ between LLM output and document text.
    
    This function is idempotent - calling it multiple times produces the same result.
    
    Args:
        text: Text with various quote characters
        
    Returns:
        Text with normalized ASCII quotes
    """
    if not text:
        return text
    
    # Normalize double quotes (curly, smart, etc.) to standard "
    # Order matters: normalize curly quotes first, then other variants
    text = text.replace('\u201C', '"')  # Left double quotation mark (U+201C) "
    text = text.replace('\u201D', '"')  # Right double quotation mark (U+201D) "
    text = text.replace('\u201E', '"')  # Double low-9 quotation mark (U+201E) „
    text = text.replace('\u00AB', '"')  # Left-pointing double angle quotation (U+00AB) «
    text = text.replace('\u00BB', '"')  # Right-pointing double angle quotation (U+00BB) »
    
    # Normalize single quotes (curly, smart, etc.) to standard '
    text = text.replace('\u2018', "'")  # Left single quotation mark (U+2018) '
    text = text.replace('\u2019', "'")  # Right single quotation mark (U+2019) '
    text = text.replace('\u201A', "'")  # Single low-9 quotation mark (U+201A) ‚
    text = text.replace('\u2039', "'")  # Single left-pointing angle quotation (U+2039) ‹
    text = text.replace('\u203A', "'")  # Single right-pointing angle quotation (U+203A) ›
    
    return text


def normalize_whitespace(text: str) -> str:
    """
    Normalize all whitespace characters to single spaces.
    
    Handles newlines, tabs, multiple spaces, and other whitespace
    that may differ between LLM output and document text.
    Also removes spaces after hyphens and punctuation that may be
    introduced by document formatting (e.g., "NON- INFRINGEMENT" -> "NON-INFRINGEMENT").
    Also handles word splits due to line breaks (e.g., "Hita chi" -> "Hitachi").
    
    Args:
        text: Text with various whitespace characters
        
    Returns:
        Text with normalized whitespace (single spaces, spaces after punctuation removed)
    """
    if not text:
        return text
    
    # Remove spaces after hyphens when they're part of compound words
    # Handles cases like "NON- INFRINGEMENT" -> "NON-INFRINGEMENT"
    # Only removes space if hyphen is NOT preceded by space (compound word, not punctuation)
    # Preserves spaces around hyphens used as punctuation (e.g., "word - word")
    # Pattern: hyphen not preceded by space, followed by whitespace
    text = re.sub(r'(?<!\s)-\s+', '-', text)
    
    # Also handle en-dash and em-dash in compound words (not preceded by space)
    text = re.sub(r'(?<!\s)[–—]\s+', lambda m: m.group(0)[0], text)
    
    # Handle specific known word splits due to line breaks
    # Only handle explicit known cases to avoid false positives
    # Common pattern: company names or technical terms split across lines
    text = re.sub(r'\bHita\s+chi\b', 'Hitachi', text, flags=re.IGNORECASE)
    
    # Handle common word splits where a single word was broken across lines
    # Pattern: short word (2-4 chars) + space + short word (2-4 chars) that form a single word
    # Only merge when both parts are very short to avoid false positives
    # Examples: "loss es" -> "losses", "non- exclusive" already handled by hyphen pattern
    common_word_splits = [
        (r'\bloss\s+es\b', 'losses'),
        (r'\bdamage\s+es\b', 'damages'),
        (r'\bexpense\s+es\b', 'expenses'),
        (r'\bjudgment\s+s\b', 'judgments'),
        (r'\bsettlement\s+s\b', 'settlements'),
        (r'\bcost\s+s\b', 'costs'),
        (r'\bfee\s+s\b', 'fees'),
        (r'\bliabilit\s+y\b', 'liability'),
        (r'\bliabilit\s+ies\b', 'liabilities'),
    ]
    for pattern, replacement in common_word_splits:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    
    # Replace all remaining whitespace (newlines, tabs, multiple spaces) with single space
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text


def normalize_subsection_markers(text: str) -> str:
    """
    Remove subsection markers that may appear in formatted documents.
    
    Handles markers like (a), (b), (i), (ii), (1), (2), etc. that may
    appear in formatted documents but not in vendor quotes extracted by LLM.
    
    Args:
        text: Text that may contain subsection markers
        
    Returns:
        Text with subsection markers removed
    """
    if not text:
        return text
    
    # Remove subsection markers at start of text or after whitespace
    # Patterns: (a), (b), (i), (ii), (iii), (iv), (v), (1), (2), etc.
    # Also handles with periods: (a.), (i.), etc.
    # Match: optional whitespace + opening paren + letter/number + optional period + closing paren + optional whitespace
    text = re.sub(r'\s*\([a-z]\)\.?\s*', ' ', text, flags=re.IGNORECASE)  # (a), (b), etc.
    text = re.sub(r'\s*\([ivxlcdm]+\)\.?\s*', ' ', text, flags=re.IGNORECASE)  # (i), (ii), (iii), (iv), (v), etc. (Roman numerals)
    text = re.sub(r'\s*\(\d+\)\.?\s*', ' ', text)  # (1), (2), (3), etc.
    
    # Clean up any double spaces created
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text


def normalize_for_matching(text: str) -> str:
    """
    Apply all normalizations to prepare text for matching.
    
    Combines: escaped quotes, quote types, whitespace normalization, subsection marker removal,
    and sentence-ending punctuation normalization.
    
    Args:
        text: Text to normalize
        
    Returns:
        Fully normalized text ready for matching
    """
    if not text:
        return text
    
    text = normalize_escaped_quotes(text)
    text = normalize_quotes(text)
    text = normalize_subsection_markers(text)  # Remove (a), (i), etc. markers
    text = normalize_whitespace(text)
    
    # Normalize sentence-ending punctuation for consistency
    # Convert semicolons at end of sentences to periods (common variation)
    # Only at the very end of the text to avoid changing mid-sentence semicolons
    text = text.rstrip()
    if text and text[-1] == ';':
        text = text[:-1] + '.'
    
    return text


def apply_exact_sentence_redlining(doc, redline_items: List[Dict[str, str]]) -> Dict[str, Any]:
    """
    Apply redlining to document with PRECISE matching.
    
    PRECISION APPROACH:
    - Only redlines EXACT vendor_quote text (with normalization fallbacks)
    - Handles multiple conflicts in the same paragraph
    - Comments are anchored to the exact runs containing the text
    - Duplicates detected by vendor_quote field only
    - Normalizes: escaped quotes, quote types, whitespace
    - Supports cross-paragraph matching for text split across paragraphs
    
    Args:
        doc: python-docx Document object
        redline_items: List of conflict items with 'text' (vendor_quote) to highlight
        
    Returns:
        Dictionary with redlining results
    """
    total_conflicts_input = len(redline_items)
    logger.info(f"APPLY_START: Processing {total_conflicts_input} conflicts across {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
    
    try:
        total_paragraphs = len(doc.paragraphs)
        failed_matches = []
        duplicates_skipped = 0
        empty_skipped = 0
        
        # Document structure logging
        logger.info(f"DOCUMENT_STRUCTURE: Total paragraphs: {total_paragraphs}")
        
        # Track seen vendor_quotes for duplicate detection (normalize for comparison)
        seen_vendor_quotes = set()
        
        def normalize_vendor_quote_for_dedup(text: str) -> str:
            """Normalize vendor_quote for duplicate detection."""
            if not text:
                return ""
            return normalize_for_matching(text).lower()
        
        # PHASE 1: Pre-scan - find all matches for all conflicts
        # This allows us to handle multiple conflicts per paragraph
        paragraph_matches = {}  # para_idx -> list of (start_pos, end_pos, comment, vendor_quote, conflict_id)
        cross_para_matches = []  # List of cross-paragraph matches to handle separately
        table_cell_matches = []  # List of table cell matches to handle separately
        
        for redline_item in redline_items:
            vendor_quote = redline_item.get('text', '').strip()
            
            if not vendor_quote:
                logger.warning("SKIP: Empty vendor_quote")
                empty_skipped += 1
                continue
            
            # Check for duplicates based on vendor_quote only
            normalized_quote = normalize_vendor_quote_for_dedup(vendor_quote)
            if normalized_quote in seen_vendor_quotes:
                logger.info(f"DUPLICATE_SKIP: vendor_quote already processed: '{vendor_quote[:50]}...'")
                duplicates_skipped += 1
                continue
            
            seen_vendor_quotes.add(normalized_quote)
            
            # Get conflict metadata
            conflict_id = redline_item.get('id', redline_item.get('clarification_id', 'Unknown'))
            serial_num = redline_item.get('serial_number', 'N/A')
            comment = redline_item.get('comment', '')
            
            logger.info(f"SCANNING: Serial={serial_num}, ID={conflict_id}, vendor_quote='{vendor_quote[:80]}...'")
            
            # Find match in document using tiered strategy
            match_result = _find_text_match(doc, vendor_quote)
            
            if match_result:
                if match_result['type'] == 'single_para':
                    para_idx = match_result['para_idx']
                    if para_idx not in paragraph_matches:
                        paragraph_matches[para_idx] = []
                    paragraph_matches[para_idx].append({
                        'start_pos': match_result['start_pos'],
                        'end_pos': match_result['end_pos'],
                        'comment': comment,
                        'vendor_quote': vendor_quote,
                        'conflict_id': conflict_id,
                        'match_type': match_result['match_type']
                    })
                    logger.info(f"FOUND: {match_result['match_type']} match in paragraph {para_idx}")
                elif match_result['type'] == 'cross_para':
                    cross_para_matches.append({
                        'paragraphs': match_result['paragraphs'],
                        'comment': comment,
                        'vendor_quote': vendor_quote,
                        'conflict_id': conflict_id
                    })
                    logger.info(f"FOUND: Cross-paragraph match spanning paragraphs {match_result['paragraphs']}")
                elif match_result['type'] == 'table_cell':
                    table_cell_matches.append({
                        'table_idx': match_result['table_idx'],
                        'row_idx': match_result['row_idx'],
                        'cell_idx': match_result['cell_idx'],
                        'comment': comment,
                        'vendor_quote': vendor_quote,
                        'conflict_id': conflict_id,
                        'match_type': match_result['match_type']
                    })
                    logger.info(f"FOUND: Table cell match in table {match_result['table_idx']}, row {match_result['row_idx']}, cell {match_result['cell_idx']}")
            else:
                failed_matches.append({
                    'text': vendor_quote,
                    'id': conflict_id,
                    'serial_number': serial_num
                })
                logger.warning(f"NO_MATCH: Serial={serial_num}, ID={conflict_id}, vendor_quote='{vendor_quote[:50]}...'")
        
        # PHASE 2: Apply redlines - process each paragraph with its matches
        matches_found = 0
        paragraphs_with_redlines = []
        
        for para_idx, matches in paragraph_matches.items():
            paragraph = doc.paragraphs[para_idx]
            para_text = paragraph.text
            
            # Sort matches by start position DESCENDING (process end to start to preserve positions)
            matches.sort(key=lambda x: x['start_pos'], reverse=True)
            
            logger.info(f"APPLYING: {len(matches)} redlines to paragraph {para_idx}")
            
            # Apply all redlines to this paragraph
            success = _apply_multiple_redlines(paragraph, para_text, matches)
            
            if success:
                matches_found += len(matches)
                paragraphs_with_redlines.append(para_idx)
        
        # PHASE 3: Handle cross-paragraph matches
        for cross_match in cross_para_matches:
            para_indices = cross_match['paragraphs']
            comment = cross_match['comment']
            
            # Apply redline to each paragraph in the cross-paragraph match
            for para_idx in para_indices:
                if para_idx < len(doc.paragraphs):
                    paragraph = doc.paragraphs[para_idx]
                    # Redline the entire paragraph text for cross-paragraph matches
                    _apply_full_paragraph_redline(paragraph, comment)
                    if para_idx not in paragraphs_with_redlines:
                        paragraphs_with_redlines.append(para_idx)
            
            matches_found += 1
            logger.info(f"CROSS_PARA_APPLIED: Redlined paragraphs {para_indices}")
        
        # PHASE 4: Handle table cell matches
        tables_with_redlines = []
        for table_match in table_cell_matches:
            table_idx = table_match['table_idx']
            row_idx = table_match['row_idx']
            cell_idx = table_match['cell_idx']
            comment = table_match['comment']
            vendor_quote = table_match['vendor_quote']
            
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    if cell_idx < len(row.cells):
                        cell = row.cells[cell_idx]
                        success = _apply_table_cell_redline(cell, vendor_quote, comment)
                        if success:
                            matches_found += 1
                            table_key = (table_idx, row_idx, cell_idx)
                            if table_key not in tables_with_redlines:
                                tables_with_redlines.append(table_key)
                            logger.info(f"TABLE_CELL_APPLIED: Redlined table {table_idx}, row {row_idx}, cell {cell_idx}")
        
        # Log results
        if paragraphs_with_redlines:
            pages_affected = set(para_idx // 20 for para_idx in paragraphs_with_redlines)
            logger.info(f"PAGE_DISTRIBUTION: {len(pages_affected)} pages affected: {sorted(pages_affected)}")
        
        total_conflicts_processed = total_conflicts_input - empty_skipped - duplicates_skipped
        success_rate = (matches_found / total_conflicts_processed * 100) if total_conflicts_processed else 0
        
        # Log comprehensive summary
        logger.info(f"REDLINE_SUMMARY: Input={total_conflicts_input}, Processed={total_conflicts_processed}, Empty skipped={empty_skipped}, Duplicates skipped={duplicates_skipped}")
        
        if failed_matches:
            logger.warning(f"REDLINING_INCOMPLETE: {len(failed_matches)} conflicts could not be matched")
            for failed in failed_matches:
                logger.warning(f"UNMATCHED: ID={failed.get('id', 'Unknown')} - '{failed.get('text', '')[:50]}...'")
        else:
            logger.info("REDLINING_SUCCESS: All conflicts successfully matched and redlined")
        
        logger.info(f"REDLINE_RESULTS: Total processed={total_conflicts_processed}, Matches={matches_found}, Failed={len(failed_matches)}, Success rate={success_rate:.1f}%")
        
        return {
            "total_paragraphs": total_paragraphs,
            "matches_found": matches_found,
            "paragraphs_with_redlines": paragraphs_with_redlines,
            "failed_matches": failed_matches,
            "pages_affected": len(set(para_idx // 20 for para_idx in paragraphs_with_redlines)) if paragraphs_with_redlines else 0
        }
        
    except Exception as e:
        logger.error(f"Error in redlining: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return {
            "total_paragraphs": 0,
            "matches_found": 0,
            "paragraphs_with_redlines": [],
            "error": str(e)
        }


def _find_text_match(doc, vendor_quote: str) -> Optional[Dict[str, Any]]:
    """
    Find vendor_quote text in document using tiered matching strategy.
    
    MATCHING TIERS:
    1. Exact match (original text)
    2. Quote-normalized match (handle ' vs " vs curly quotes)
    3. Quote + whitespace normalized match (handle quote and whitespace variations)
    4. Fully normalized match (quotes + whitespace + case-insensitive)
    5. Cross-paragraph match (text spans multiple paragraphs)
    6. Table cell match (searches within table cells)
    
    Args:
        doc: python-docx Document object
        vendor_quote: Text to find in document
        
    Returns:
        Match result dict or None if not found
    """
    # Prepare normalized versions
    # Note: vendor_quote may already be normalized from parsing, so we normalize again (idempotent)
    quote_normalized = normalize_quotes(normalize_escaped_quotes(vendor_quote))
    quote_ws_normalized = normalize_whitespace(normalize_subsection_markers(quote_normalized))
    fully_normalized = normalize_for_matching(vendor_quote)
    
    # First, search through all tables (for table-formatted exception documents)
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                
                # Normalize cell text for matching (same normalization as vendor_quote)
                cell_quote_normalized = normalize_quotes(cell_text)
                cell_quote_ws_normalized = normalize_whitespace(normalize_subsection_markers(cell_quote_normalized))
                cell_fully_normalized = normalize_for_matching(cell_text).lower()
                
                # Check all matching tiers for table cells
                if vendor_quote in cell_text:
                    logger.info(f"MATCH_FOUND: TIER 1 (exact) in table {table_idx}, row {row_idx}, cell {cell_idx}")
                    return {
                        'type': 'table_cell',
                        'table_idx': table_idx,
                        'row_idx': row_idx,
                        'cell_idx': cell_idx,
                        'match_type': 'exact'
                    }
                
                if quote_normalized in cell_quote_normalized:
                    logger.info(f"MATCH_FOUND: TIER 2 (quote_normalized) in table {table_idx}, row {row_idx}, cell {cell_idx}")
                    return {
                        'type': 'table_cell',
                        'table_idx': table_idx,
                        'row_idx': row_idx,
                        'cell_idx': cell_idx,
                        'match_type': 'quote_normalized'
                    }
                
                if quote_ws_normalized in cell_quote_ws_normalized:
                    logger.info(f"MATCH_FOUND: TIER 3 (quote_whitespace_normalized) in table {table_idx}, row {row_idx}, cell {cell_idx}")
                    return {
                        'type': 'table_cell',
                        'table_idx': table_idx,
                        'row_idx': row_idx,
                        'cell_idx': cell_idx,
                        'match_type': 'quote_whitespace_normalized'
                    }
                
                if quote_ws_normalized.lower() in cell_quote_ws_normalized.lower():
                    logger.info(f"MATCH_FOUND: TIER 3b (quote_whitespace_normalized_case_insensitive) in table {table_idx}, row {row_idx}, cell {cell_idx}")
                    return {
                        'type': 'table_cell',
                        'table_idx': table_idx,
                        'row_idx': row_idx,
                        'cell_idx': cell_idx,
                        'match_type': 'quote_whitespace_normalized_case_insensitive'
                    }
                
                if fully_normalized.lower() in cell_fully_normalized:
                    logger.info(f"MATCH_FOUND: TIER 4 (fully_normalized) in table {table_idx}, row {row_idx}, cell {cell_idx}")
                    return {
                        'type': 'table_cell',
                        'table_idx': table_idx,
                        'row_idx': row_idx,
                        'cell_idx': cell_idx,
                        'match_type': 'fully_normalized'
                    }
                
                # TIER 4b: Check if vendor_quote is a prefix/substring of cell text (for table exceptions)
                # This handles cases where Claude extracts partial text or text with formatting differences
                if len(vendor_quote) > 50:  # Only for reasonably long quotes
                    # Check if vendor quote is prefix of cell or cell is prefix of vendor quote (bidirectional)
                    if cell_quote_ws_normalized.startswith(quote_ws_normalized) or quote_ws_normalized.startswith(cell_quote_ws_normalized):
                        logger.info(f"MATCH_FOUND: TIER 4b (prefix/substring) in table {table_idx}, row {row_idx}, cell {cell_idx}")
                        return {
                            'type': 'table_cell',
                            'table_idx': table_idx,
                            'row_idx': row_idx,
                            'cell_idx': cell_idx,
                            'match_type': 'prefix_substring'
                        }
                    elif cell_fully_normalized.startswith(fully_normalized.lower()) or fully_normalized.lower().startswith(cell_fully_normalized):
                        logger.info(f"MATCH_FOUND: TIER 4c (fully_normalized prefix/substring) in table {table_idx}, row {row_idx}, cell {cell_idx}")
                        return {
                            'type': 'table_cell',
                            'table_idx': table_idx,
                            'row_idx': row_idx,
                            'cell_idx': cell_idx,
                            'match_type': 'fully_normalized_prefix_substring'
                        }
                    # Also check if vendor quote appears as substring within cell (not just prefix)
                    elif quote_ws_normalized in cell_quote_ws_normalized or cell_quote_ws_normalized in quote_ws_normalized:
                        logger.info(f"MATCH_FOUND: TIER 4d (substring within cell) in table {table_idx}, row {row_idx}, cell {cell_idx}")
                        return {
                            'type': 'table_cell',
                            'table_idx': table_idx,
                            'row_idx': row_idx,
                            'cell_idx': cell_idx,
                            'match_type': 'substring_within_cell'
                        }
    
    # Search through all paragraphs (fallback if not found in tables)
    paragraphs_checked = 0
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text
        if not para_text.strip():
            continue
        
        paragraphs_checked += 1
        
        # TIER 1: Exact match
        if vendor_quote in para_text:
            start_pos = para_text.find(vendor_quote)
            logger.info(f"MATCH_FOUND: TIER 1 (exact) in paragraph {para_idx} at position {start_pos}")
            return {
                'type': 'single_para',
                'para_idx': para_idx,
                'start_pos': start_pos,
                'end_pos': start_pos + len(vendor_quote),
                'match_type': 'exact'
            }
        
        # TIER 2: Quote-normalized match (quotes only, preserve whitespace)
        para_quote_normalized = normalize_quotes(para_text)
        if quote_normalized in para_quote_normalized:
            start_pos = para_quote_normalized.find(quote_normalized)
            logger.info(f"MATCH_FOUND: TIER 2 (quote_normalized) in paragraph {para_idx} at position {start_pos}")
            # Map back to original positions (approximate - quotes are same length)
            return {
                'type': 'single_para',
                'para_idx': para_idx,
                'start_pos': start_pos,
                'end_pos': start_pos + len(vendor_quote),
                'match_type': 'quote_normalized'
            }
        
        # TIER 3: Quote + whitespace normalized match
        para_quotes_only = normalize_quotes(para_text)
        para_quote_ws_normalized = normalize_whitespace(normalize_subsection_markers(para_quotes_only))
        if quote_ws_normalized in para_quote_ws_normalized:
            # Find position in normalized text, then map back
            norm_start = para_quote_ws_normalized.find(quote_ws_normalized)
            logger.info(f"MATCH_FOUND: TIER 3 (quote_whitespace_normalized) in paragraph {para_idx} at normalized position {norm_start}")
            # Map normalized position back to original position
            start_pos = _map_normalized_to_original_position(para_text, norm_start)
            end_pos = _map_normalized_to_original_position(para_text, norm_start + len(quote_ws_normalized))
            return {
                'type': 'single_para',
                'para_idx': para_idx,
                'start_pos': start_pos,
                'end_pos': end_pos,
                'match_type': 'quote_whitespace_normalized'
            }
        
        # TIER 3b: Quote + whitespace normalized + case-insensitive match (for case differences)
        if quote_ws_normalized.lower() in para_quote_ws_normalized.lower():
            norm_start = para_quote_ws_normalized.lower().find(quote_ws_normalized.lower())
            logger.info(f"MATCH_FOUND: TIER 3b (quote_whitespace_normalized_case_insensitive) in paragraph {para_idx} at normalized position {norm_start}")
            start_pos = _map_normalized_to_original_position(para_text, norm_start)
            end_pos = _map_normalized_to_original_position(para_text, norm_start + len(quote_ws_normalized))
            return {
                'type': 'single_para',
                'para_idx': para_idx,
                'start_pos': start_pos,
                'end_pos': end_pos,
                'match_type': 'quote_whitespace_normalized_case_insensitive'
            }
        
        # TIER 4: Fully normalized + case-insensitive match
        para_fully_normalized = normalize_for_matching(para_text).lower()
        if fully_normalized.lower() in para_fully_normalized:
            norm_start = para_fully_normalized.find(fully_normalized.lower())
            logger.info(f"MATCH_FOUND: TIER 4 (fully_normalized) in paragraph {para_idx} at normalized position {norm_start}")
            start_pos = _map_normalized_to_original_position(para_text, norm_start)
            end_pos = _map_normalized_to_original_position(para_text, norm_start + len(fully_normalized))
            return {
                'type': 'single_para',
                'para_idx': para_idx,
                'start_pos': start_pos,
                'end_pos': end_pos,
                'match_type': 'fully_normalized'
            }
    
    # TIER 5: Partial/truncated quote matching
    # If vendor quote appears to be truncated (ends mid-sentence or is suspiciously short),
    # or is very long (might span paragraphs), try to match it as a prefix/substring
    partial_match = _find_partial_match(doc, vendor_quote, quote_ws_normalized, fully_normalized, force_substring_match=False)
    if partial_match:
        logger.info(f"MATCH_FOUND: TIER 5 (partial/truncated) in paragraph {partial_match.get('para_idx')}")
        return partial_match
    
    # TIER 5b: Substring matching fallback for non-truncated quotes
    # If exact match failed but quote is long or we suspect formatting differences,
    # try substring matching as a last resort
    if len(vendor_quote) > 200:  # Only for reasonably long quotes to avoid false positives
        substring_match = _find_partial_match(doc, vendor_quote, quote_ws_normalized, fully_normalized, force_substring_match=True)
        if substring_match:
            logger.info(f"MATCH_FOUND: TIER 5b (substring fallback) in paragraph {substring_match.get('para_idx')}")
            return substring_match
    
    # TIER 6: Cross-paragraph matching (enhanced for long quotes)
    cross_match = _find_cross_paragraph_match(doc, vendor_quote, fully_normalized, quote_ws_normalized, fully_normalized)
    if cross_match:
        logger.info(f"MATCH_FOUND: TIER 6 (cross_paragraph) across paragraphs {cross_match.get('paragraphs', [])}")
        return cross_match
    
    logger.warning(f"MATCH_FAILED: Could not find vendor_quote in document. vendor_quote='{vendor_quote[:200]}...'")
    logger.warning(f"MATCH_FAILED: Checked {len([p for p in doc.paragraphs if p.text.strip()])} non-empty paragraphs")
    return None


def _map_normalized_to_original_position(original_text: str, normalized_pos: int) -> int:
    """
    Map a position in normalized text back to approximate position in original text.
    
    This is an approximation since whitespace normalization changes positions.
    
    Args:
        original_text: Original text before normalization
        normalized_pos: Position in normalized text
        
    Returns:
        Approximate position in original text
    """
    if normalized_pos <= 0:
        return 0
    
    # Walk through original text, counting non-collapsed characters
    normalized_count = 0
    in_whitespace = False
    
    for i, char in enumerate(original_text):
        if char in ' \t\n\r':
            if not in_whitespace:
                normalized_count += 1  # Collapsed whitespace counts as 1
                in_whitespace = True
        else:
            normalized_count += 1
            in_whitespace = False
        
        if normalized_count >= normalized_pos:
            return i + 1
    
    return len(original_text)


def _find_partial_match(doc, vendor_quote: str, quote_ws_normalized: str, fully_normalized: str, force_substring_match: bool = False) -> Optional[Dict[str, Any]]:
    """
    Find vendor_quote using partial/substring matching.
    
    If vendor quote is a prefix or substring of document text (after normalization), treat it as a match.
    This handles cases where:
    1. The LLM extracted an incomplete quote (truncated)
    2. Exact match failed but quote exists as substring (quote/formatting differences)
    3. Long quotes that span multiple paragraphs
    
    Args:
        doc: python-docx Document object
        vendor_quote: Original vendor quote text
        quote_ws_normalized: Quote and whitespace normalized vendor quote
        fully_normalized: Fully normalized vendor quote
        force_substring_match: If True, attempt substring matching even for non-truncated quotes
        
    Returns:
        Match result dict or None if not found
    """
    # Check if vendor quote appears truncated (ends mid-sentence, suspiciously short, etc.)
    is_likely_truncated = (
        len(vendor_quote) < 100 or  # Suspiciously short
        not vendor_quote.rstrip().endswith(('.', '!', '?', ';', ':', ')', ']', '}')) or  # Doesn't end with punctuation
        vendor_quote.rstrip().endswith(('or', 'and', 'the', 'a', 'an', 'to', 'of', 'in', 'for', 'with'))  # Ends with common words
    )
    
    # For long quotes (>500 chars), also try substring matching as they might span paragraphs
    is_long_quote = len(vendor_quote) > 500
    
    if not (is_likely_truncated or force_substring_match or is_long_quote):
        return None
    
    if is_likely_truncated:
        logger.info(f"MATCH_PARTIAL_CHECK: vendor_quote appears truncated (length={len(vendor_quote)}, ends_with='{vendor_quote[-20:]}')")
    elif is_long_quote:
        logger.info(f"MATCH_PARTIAL_CHECK: vendor_quote is long ({len(vendor_quote)} chars), attempting substring matching")
    elif force_substring_match:
        logger.info(f"MATCH_PARTIAL_CHECK: forcing substring match attempt (exact match failed)")
    
    # Search through paragraphs for prefix match
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text
        if not para_text.strip():
            continue
        
        # Normalize paragraph text
        para_quotes_only = normalize_quotes(para_text)
        para_quote_ws_normalized = normalize_whitespace(normalize_subsection_markers(para_quotes_only))
        para_fully_normalized = normalize_for_matching(para_text).lower()
        
        # Check if vendor quote is a prefix of paragraph (after normalization)
        # Try multiple normalization levels
        if para_quote_ws_normalized.startswith(quote_ws_normalized):
            # Found prefix match - vendor quote is start of paragraph
            logger.info(f"MATCH_PARTIAL_FOUND: vendor_quote is prefix of paragraph {para_idx} (quote_ws_normalized)")
            start_pos = 0
            # Map normalized length back to original position
            end_pos = _map_normalized_to_original_position(para_text, len(quote_ws_normalized))
            return {
                'type': 'single_para',
                'para_idx': para_idx,
                'start_pos': start_pos,
                'end_pos': end_pos,
                'match_type': 'partial_truncated',
                'is_truncated': True
            }
        elif para_quote_ws_normalized.lower().startswith(quote_ws_normalized.lower()):
            # Found prefix match with case-insensitive check
            logger.info(f"MATCH_PARTIAL_FOUND: vendor_quote is prefix of paragraph {para_idx} (quote_ws_normalized_case_insensitive)")
            start_pos = 0
            end_pos = _map_normalized_to_original_position(para_text, len(quote_ws_normalized))
            return {
                'type': 'single_para',
                'para_idx': para_idx,
                'start_pos': start_pos,
                'end_pos': end_pos,
                'match_type': 'partial_truncated_case_insensitive',
                'is_truncated': True
            }
        elif para_fully_normalized.startswith(fully_normalized.lower()):
            # Found prefix match with full normalization
            logger.info(f"MATCH_PARTIAL_FOUND: vendor_quote is prefix of paragraph {para_idx} (fully_normalized)")
            start_pos = 0
            end_pos = _map_normalized_to_original_position(para_text, len(fully_normalized))
            return {
                'type': 'single_para',
                'para_idx': para_idx,
                'start_pos': start_pos,
                'end_pos': end_pos,
                'match_type': 'partial_truncated',
                'is_truncated': True
            }
        
        # Check if vendor quote appears within paragraph (not just at start)
        # This handles cases where vendor quote is a substring (e.g., starts after " Conflicts. ")
        # Try multiple normalization levels - check this for ALL paragraphs, not just when not a prefix
        if quote_ws_normalized in para_quote_ws_normalized:
            norm_start = para_quote_ws_normalized.find(quote_ws_normalized)
            logger.info(f"MATCH_PARTIAL_SUBSTRING_CHECK: paragraph {para_idx}, quote_ws_normalized found at position {norm_start}")
            # Only return if it's not already a prefix match (avoid duplicate)
            if norm_start > 0:
                logger.info(f"MATCH_PARTIAL_FOUND: vendor_quote found within paragraph {para_idx} at normalized position {norm_start} (quote_ws_normalized)")
                start_pos = _map_normalized_to_original_position(para_text, norm_start)
                end_pos = _map_normalized_to_original_position(para_text, norm_start + len(quote_ws_normalized))
                return {
                    'type': 'single_para',
                    'para_idx': para_idx,
                    'start_pos': start_pos,
                    'end_pos': end_pos,
                    'match_type': 'partial_substring',
                    'is_truncated': is_likely_truncated
                }
        elif quote_ws_normalized.lower() in para_quote_ws_normalized.lower():
            # Case-insensitive substring check
            norm_start = para_quote_ws_normalized.lower().find(quote_ws_normalized.lower())
            if norm_start > 0:
                logger.info(f"MATCH_PARTIAL_FOUND: vendor_quote found within paragraph {para_idx} at normalized position {norm_start} (quote_ws_normalized_case_insensitive)")
                start_pos = _map_normalized_to_original_position(para_text, norm_start)
                end_pos = _map_normalized_to_original_position(para_text, norm_start + len(quote_ws_normalized))
                return {
                    'type': 'single_para',
                    'para_idx': para_idx,
                    'start_pos': start_pos,
                    'end_pos': end_pos,
                    'match_type': 'partial_substring_case_insensitive',
                    'is_truncated': is_likely_truncated
                }
        elif fully_normalized.lower() in para_fully_normalized:
            norm_start = para_fully_normalized.find(fully_normalized.lower())
            # Only return if it's not already a prefix match (avoid duplicate)
            if norm_start > 0:
                logger.info(f"MATCH_PARTIAL_FOUND: vendor_quote found within paragraph {para_idx} at normalized position {norm_start} (fully_normalized)")
                start_pos = _map_normalized_to_original_position(para_text, norm_start)
                end_pos = _map_normalized_to_original_position(para_text, norm_start + len(fully_normalized))
                return {
                    'type': 'single_para',
                    'para_idx': para_idx,
                    'start_pos': start_pos,
                    'end_pos': end_pos,
                    'match_type': 'partial_substring',
                    'is_truncated': is_likely_truncated
                }
        # Log divergence details only when first 50 chars match but full quote doesn't (for debugging)
        elif len(quote_ws_normalized) > 50 and quote_ws_normalized[:50] in para_quote_ws_normalized:
            pos50 = para_quote_ws_normalized.find(quote_ws_normalized[:50])
            if pos50 >= 0:
                # Find where text diverges
                for i in range(50, min(len(quote_ws_normalized), len(para_quote_ws_normalized) - pos50)):
                    if quote_ws_normalized[i] != para_quote_ws_normalized[pos50 + i]:
                        logger.warning(f"MATCH_PARTIAL_SUBSTRING_CHECK: First 50 chars match in paragraph {para_idx}, but diverges at position {i}. Quote char: '{quote_ws_normalized[i]}' (U+{ord(quote_ws_normalized[i]):04X}), Doc char: '{para_quote_ws_normalized[pos50 + i]}' (U+{ord(para_quote_ws_normalized[pos50 + i]):04X})")
                        break
                else:
                    if len(quote_ws_normalized) > len(para_quote_ws_normalized) - pos50:
                        logger.warning(f"MATCH_PARTIAL_SUBSTRING_CHECK: Quote is longer than available text in paragraph {para_idx}. Quote length={len(quote_ws_normalized)}, Available={len(para_quote_ws_normalized) - pos50}")
    
    return None


def _find_cross_paragraph_match(doc, vendor_quote: str, normalized_quote: str, quote_ws_normalized: str = None, fully_normalized: str = None) -> Optional[Dict[str, Any]]:
    """
    Find vendor_quote that spans multiple paragraphs.
    
    Joins consecutive paragraphs and searches for the text.
    Enhanced for long quotes with better normalization handling.
    
    Args:
        doc: python-docx Document object
        vendor_quote: Original text to find
        normalized_quote: Normalized version of text
        
    Returns:
        Match result dict or None if not found
    """
    paragraphs = doc.paragraphs
    
    # Determine window size based on quote length
    # Long quotes (>500 chars) might span more paragraphs
    # Increased window sizes to handle quotes that span many paragraphs
    if len(vendor_quote) > 1000:
        window_sizes = [5, 6, 7, 8, 10]  # Try larger windows for very long quotes
    elif len(vendor_quote) > 500:
        window_sizes = [4, 5, 6, 7, 8]  # Increased from [3, 4, 5]
    elif len(vendor_quote) > 300:
        window_sizes = [3, 4, 5, 6]  # Added intermediate tier
    else:
        window_sizes = [2, 3, 4]  # Added 4 for medium quotes
    
    # Try joining consecutive paragraphs with different window sizes
    for window_size in window_sizes:
        for start_idx in range(len(paragraphs) - window_size + 1):
            # Join paragraphs with space
            joined_paras = []
            joined_text = ""
            for i in range(window_size):
                para_text = paragraphs[start_idx + i].text
                if para_text.strip():
                    joined_paras.append(start_idx + i)
                    joined_text += para_text + " "
            
            if not joined_text.strip():
                continue
            
            joined_text = joined_text.strip()
            
            # Try multiple normalization levels for better matching
            # First try quote+whitespace normalized if available
            if quote_ws_normalized:
                joined_quote_ws = normalize_whitespace(normalize_subsection_markers(normalize_quotes(joined_text)))
                if quote_ws_normalized in joined_quote_ws:
                    logger.info(f"CROSS_PARA_MATCH: Found (quote_ws_normalized) in paragraphs {joined_paras}")
                    return {
                        'type': 'cross_para',
                        'paragraphs': joined_paras,
                        'match_type': 'cross_paragraph_quote_ws'
                    }
                # Also try case-insensitive version
                if quote_ws_normalized.lower() in joined_quote_ws.lower():
                    logger.info(f"CROSS_PARA_MATCH: Found (quote_ws_normalized_case_insensitive) in paragraphs {joined_paras}")
                    return {
                        'type': 'cross_para',
                        'paragraphs': joined_paras,
                        'match_type': 'cross_paragraph_quote_ws_case_insensitive'
                    }
            
            # Try fully normalized
            joined_normalized = normalize_for_matching(joined_text).lower()
            if normalized_quote.lower() in joined_normalized:
                logger.info(f"CROSS_PARA_MATCH: Found (fully_normalized) in paragraphs {joined_paras}")
                return {
                    'type': 'cross_para',
                    'paragraphs': joined_paras,
                    'match_type': 'cross_paragraph'
                }
            
            # Also try with fully_normalized if available
            if fully_normalized:
                if fully_normalized.lower() in joined_normalized:
                    logger.info(f"CROSS_PARA_MATCH: Found (fully_normalized) in paragraphs {joined_paras}")
                    return {
                        'type': 'cross_para',
                        'paragraphs': joined_paras,
                        'match_type': 'cross_paragraph_fully'
                    }
    
    return None


def _apply_multiple_redlines(paragraph, para_text: str, matches: List[Dict]) -> bool:
    """
    Apply multiple redlines to a single paragraph.
    
    Processes matches from end to start to preserve character positions.
    
    Args:
        paragraph: python-docx Paragraph object
        para_text: Original paragraph text
        matches: List of match dicts sorted by start_pos descending
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Build list of text segments with formatting info
        # Start with the full text, then split at each redline point
        segments = []  # List of (text, is_redline, comment)
        
        current_pos = len(para_text)
        
        for match in matches:
            start_pos = match['start_pos']
            end_pos = match['end_pos']
            comment = match['comment']
            
            # Bounds checking
            start_pos = max(0, min(start_pos, len(para_text)))
            end_pos = max(start_pos, min(end_pos, len(para_text)))
            
            # Add text after this redline (if any)
            if current_pos > end_pos:
                after_text = para_text[end_pos:current_pos]
                if after_text:
                    segments.insert(0, {'text': after_text, 'is_redline': False, 'comment': None})
            
            # Add the redlined text
            target_text = para_text[start_pos:end_pos]
            if target_text:
                segments.insert(0, {'text': target_text, 'is_redline': True, 'comment': comment})
            
            current_pos = start_pos
        
        # Add any remaining text at the beginning
        if current_pos > 0:
            before_text = para_text[:current_pos]
            if before_text:
                segments.insert(0, {'text': before_text, 'is_redline': False, 'comment': None})
        
        # Clear existing runs
        for run in paragraph.runs:
            run.clear()
        
        # Rebuild paragraph with all segments
        for segment in segments:
            run = paragraph.add_run(segment['text'])
            
            if segment['is_redline']:
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                run.font.strike = True  # Strikethrough
                
                # Add comment if present
                if segment['comment'] and segment['comment'].strip():
                    try:
                        run.add_comment(segment['comment'], author="One L", initials="1L")
                        logger.info(f"COMMENT_ADDED: '{segment['comment'][:50]}...' on text '{segment['text'][:30]}...'")
                    except Exception as comment_err:
                        logger.warning(f"COMMENT_FAILED: {comment_err}")
        
        logger.info(f"MULTI_REDLINE_APPLIED: {len(matches)} redlines in paragraph")
        return True
        
    except Exception as e:
        logger.error(f"Error applying multiple redlines: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return False


def _apply_full_paragraph_redline(paragraph, comment: str):
    """
    Apply redline to entire paragraph (for cross-paragraph matches).
    
    Args:
        paragraph: python-docx Paragraph object
        comment: Comment to attach
    """
    try:
        para_text = paragraph.text
        if not para_text.strip():
            return
        
        # Clear existing runs
        for run in paragraph.runs:
            run.clear()
        
        # Add redlined text
        redline_run = paragraph.add_run(para_text)
        redline_run.font.color.rgb = RGBColor(255, 0, 0)
        redline_run.font.strike = True
        
        # Add comment (only to first paragraph in cross-para match)
        if comment and comment.strip():
            try:
                redline_run.add_comment(comment, author="One L", initials="1L")
            except Exception:
                pass
        
        logger.info(f"FULL_PARA_REDLINE: Applied to '{para_text[:50]}...'")
        
    except Exception as e:
        logger.error(f"Error applying full paragraph redline: {str(e)}")


def _apply_table_cell_redline(cell, vendor_quote: str, comment: str) -> bool:
    """
    Apply redlining to a table cell containing vendor exception text.
    
    Args:
        cell: python-docx Table cell object
        vendor_quote: The vendor exception text to redline
        comment: Comment to attach
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Get all paragraphs in the cell
        cell_paragraphs = cell.paragraphs
        if not cell_paragraphs:
            return False
        
        # Search through all paragraphs in the cell to find the vendor_quote
        quote_normalized = normalize_quotes(normalize_escaped_quotes(vendor_quote))
        quote_ws_normalized = normalize_whitespace(normalize_subsection_markers(quote_normalized))
        fully_normalized = normalize_for_matching(vendor_quote)
        
        for para in cell_paragraphs:
            para_text = para.text
            if not para_text.strip():
                continue
            
            # Try to find the vendor_quote in this paragraph
            found_match = False
            match_type = None
            
            if vendor_quote in para_text:
                found_match = True
                match_type = 'exact'
            elif quote_normalized in normalize_quotes(para_text):
                found_match = True
                match_type = 'quote_normalized'
            elif quote_ws_normalized in normalize_whitespace(normalize_subsection_markers(normalize_quotes(para_text))):
                found_match = True
                match_type = 'quote_whitespace_normalized'
            elif fully_normalized.lower() in normalize_for_matching(para_text).lower():
                found_match = True
                match_type = 'fully_normalized'
            
            if found_match:
                # Apply redlining to the entire paragraph (since cell text may be formatted)
                # Clear existing runs
                for run in para.runs:
                    run.clear()
                
                # Add redlined text
                redline_run = para.add_run(para_text)
                redline_run.font.color.rgb = RGBColor(255, 0, 0)
                redline_run.font.strike = True
                
                # Add comment
                if comment and comment.strip():
                    try:
                        redline_run.add_comment(comment, author="One L", initials="1L")
                        logger.info(f"TABLE_CELL_COMMENT: Added comment to cell")
                    except Exception as comment_err:
                        logger.warning(f"TABLE_CELL_COMMENT_FAILED: {comment_err}")
                
                logger.info(f"TABLE_CELL_REDLINE: Applied {match_type} match to cell paragraph")
                return True
        
        # If we didn't find an exact match, redline the entire cell content
        # This handles cases where the text might be split across paragraphs or formatted differently
        if cell_paragraphs:
            first_para = cell_paragraphs[0]
            cell_text = cell.text.strip()
            if cell_text:
                # Clear all paragraphs
                for para in cell_paragraphs:
                    for run in para.runs:
                        run.clear()
                
                # Add redlined text to first paragraph
                redline_run = first_para.add_run(cell_text)
                redline_run.font.color.rgb = RGBColor(255, 0, 0)
                redline_run.font.strike = True
                
                # Add comment
                if comment and comment.strip():
                    try:
                        redline_run.add_comment(comment, author="One L", initials="1L")
                    except Exception:
                        pass
                
                logger.info(f"TABLE_CELL_REDLINE: Applied to entire cell content")
                return True
        
        return False
        
    except Exception as e:
        logger.error(f"Error applying table cell redline: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return False


def _get_bucket_name(bucket_type: str) -> str:
    """Get the appropriate bucket name based on type."""
    if bucket_type == "knowledge":
        return os.environ.get("KNOWLEDGE_BUCKET")
    elif bucket_type == "user_documents":
        return os.environ.get("USER_DOCUMENTS_BUCKET")
    elif bucket_type == "agent_processing":
        return os.environ.get("AGENT_PROCESSING_BUCKET")
    else:
        raise ValueError(f"Invalid bucket_type: {bucket_type}")


def _copy_document_to_processing(original_s3_key: str, source_bucket: str, agent_bucket: str) -> str:
    """Copy document to agent processing bucket with organized structure."""
    
    from datetime import datetime
    import uuid
    
    try:
        # Create organized folder structure
        timestamp = datetime.utcnow().strftime("%Y/%m/%d")
        unique_id = str(uuid.uuid4())[:8]
        
        # Extract filename
        filename = original_s3_key.split('/')[-1]
        
        # Create new key with organized structure
        agent_key = f"input/{timestamp}/{unique_id}_{filename}"
        
        # Copy object to agent processing bucket
        s3_client.copy_object(
            CopySource={'Bucket': source_bucket, 'Key': original_s3_key},
            Bucket=agent_bucket,
            Key=agent_key,
            MetadataDirective='COPY'
        )
        

        return agent_key
        
    except Exception as e:
        logger.error(f"Error copying document to processing bucket: {str(e)}")
        raise


def _download_and_load_document(bucket: str, s3_key: str):
    """
    Download document from S3 and load it using python-docx.
    
    Args:
        bucket: S3 bucket name
        s3_key: S3 key of the document
        
    Returns:
        python-docx Document object
    """
    
    try:

        
        # Download the document from S3
        response = s3_client.get_object(Bucket=bucket, Key=s3_key)
        document_content = response['Body'].read()
        
        # Load the document using python-docx
        doc = Document(io.BytesIO(document_content))
        

        return doc
        
    except Exception as e:
        logger.error(f"Error downloading and loading document: {str(e)}")
        raise Exception(f"Failed to download and load document: {str(e)}")


def _create_redlined_filename(original_s3_key: str, session_id: str = None, user_id: str = None) -> str:
    """
    Create a filename for the redlined document.
    
    Args:
        original_s3_key: Original document S3 key
        session_id: Session ID for organizing output files
        user_id: User ID for organizing output files
        
    Returns:
        New S3 key for redlined document
    """
    
    # Split the path and filename
    path_parts = original_s3_key.split('/')
    filename = path_parts[-1]
    
    # Add redlined suffix
    name_parts = filename.rsplit('.', 1)
    if len(name_parts) == 2:
        redlined_filename = f"{name_parts[0]}_REDLINED.{name_parts[1]}"
    else:
        redlined_filename = f"{filename}_REDLINED"
    
    # Create session-based path if session info is provided
    if session_id and user_id:
        redlined_s3_key = f"sessions/{user_id}/{session_id}/output/{redlined_filename}"
    else:
        # Fallback to original logic for backwards compatibility
        if len(path_parts) > 1:
            # Change 'input' to 'output' for redlined documents
            new_path_parts = [part if part != 'input' else 'output' for part in path_parts[:-1]]
            redlined_s3_key = '/'.join(new_path_parts) + '/' + redlined_filename
        else:
            redlined_s3_key = f"output/{redlined_filename}"
    
    return redlined_s3_key




def save_analysis_to_dynamodb(
    analysis_id: str,
    document_s3_key: str,
    analysis_data: str,
    bucket_type: str,
    usage_data: Dict[str, Any],
    thinking: str = "",
    citations: List[Dict[str, Any]] = None,
    session_id: str = None,
    user_id: str = None,
    redlined_result: Dict[str, Any] = None,
    timestamp: str = None
) -> Dict[str, Any]:
    """
    Save analysis results to DynamoDB table including parsed conflicts.
    
    Args:
        analysis_id: Unique identifier for this analysis
        document_s3_key: S3 key of the analyzed document
        analysis_data: The analysis text containing conflicts table
        bucket_type: Source bucket type
        usage_data: Model usage statistics
        thinking: AI thinking process
        citations: Knowledge base citations
        session_id: Optional session ID for linking
        user_id: Optional user ID for linking
        redlined_result: Optional redlined document result
        timestamp: Optional timestamp (if provided, updates existing job record; if None, creates new record)
        
    Returns:
        Dictionary indicating success/failure of save operation
    """
    
    try:
        table_name = os.environ.get('ANALYSIS_TABLE')
        if not table_name:
            return {
                "success": False,
                "error": "ANALYSIS_TABLE environment variable not set"
            }
        
        import boto3
        from datetime import datetime
        dynamodb_resource = boto3.resource('dynamodb')
        table = dynamodb_resource.Table(table_name)
        # Use provided timestamp if available (for updating existing job), otherwise create new one
        if timestamp is None:
            timestamp = datetime.utcnow().isoformat()
        
        # Parse conflicts from analysis data for structured storage
        redline_items = parse_conflicts_for_redlining(analysis_data)
        
        # Convert redline format to DynamoDB format with better naming
        conflicts = []
        for item in redline_items:
            # Extract rationale - prefer direct field, fallback to parsing comment
            rationale = item.get('rationale', '')
            if not rationale and 'comment' in item:
                # Parse rationale from comment format: "CONFLICT ID (type): rationale\n\nReference: ..."
                comment = item['comment']
                if '): ' in comment:
                    # Extract everything after "): " and before "\n\nReference:" if present
                    rationale_part = comment.split('): ', 1)[-1]
                    if '\n\nReference:' in rationale_part:
                        rationale = rationale_part.split('\n\nReference:')[0].strip()
                    else:
                        rationale = rationale_part.strip()
                else:
                    rationale = comment
            
            conflicts.append({
                'clarification_id': item.get('clarification_id', ''),
                'vendor_conflict': item.get('text', ''),  # Exact text from vendor document (better naming)
                'summary': item.get('summary', ''),  # 20-40 word context from JSON
                'source_doc': item.get('source_doc', ''),
                'clause_ref': item.get('clause_ref', 'N/A'),
                'conflict_type': item.get('conflict_type', ''),
                'rationale': rationale or item.get('comment', '')  # Fallback to full comment if no rationale
            })
        
        # Validate and normalize JSON analysis_data before storing
        normalized_analysis_data = None
        if analysis_data:
            import json
            import re
            
            # Try to extract and validate JSON from analysis_data
            json_match = re.search(r'\[[\s\S]*\]', analysis_data)
            if json_match:
                try:
                    json_str = json_match.group(0)
                    # Validate JSON by parsing it
                    parsed_json = json.loads(json_str)
                    # Re-serialize to ensure clean, normalized JSON
                    normalized_analysis_data = json.dumps(parsed_json, indent=2, ensure_ascii=False)
                    logger.info(f"DynamoDB: Storing validated JSON with {len(parsed_json)} conflicts")
                except json.JSONDecodeError as e:
                    logger.warning(f"DynamoDB: Invalid JSON in analysis_data, storing as-is: {e}")
                    normalized_analysis_data = analysis_data
            else:
                # No JSON found, store as-is (backwards compatibility with markdown)
                normalized_analysis_data = analysis_data
                logger.info("DynamoDB: No JSON found in analysis_data, storing as text")
        
        # Prepare streamlined item for DynamoDB - focusing only on conflicts data
        item = {
            'analysis_id': analysis_id,
            'timestamp': timestamp,
            'document_s3_key': document_s3_key,
            'bucket_type': bucket_type,
            'conflicts_count': len(conflicts),
            'conflicts': conflicts
        }

        if normalized_analysis_data:
            item['analysis_data'] = normalized_analysis_data

        if usage_data:
            item['usage'] = usage_data

        if thinking:
            item['thinking'] = thinking

        if citations:
            item['citations'] = citations
        
        # Add session and user linking if provided
        if session_id:
            item['session_id'] = session_id
        if user_id:
            item['user_id'] = user_id

        if redlined_result:
            redlined_key = redlined_result.get('redlined_document')
            if redlined_key:
                item['redlined_document_s3_key'] = redlined_key
        
        # CRITICAL: If timestamp is provided, we're updating an existing job record
        # Preserve existing fields by reading the existing record first and merging
        if timestamp:
            try:
                existing_item = table.get_item(
                    Key={
                        'analysis_id': analysis_id,
                        'timestamp': timestamp
                    }
                ).get('Item')
                
                if existing_item:
                    # Merge: preserve all existing fields, then update with new data
                    # This ensures we don't lose fields like execution_arn, created_at, etc.
                    merged_item = existing_item.copy()
                    merged_item.update(item)  # New data overrides existing
                    # Always update updated_at timestamp
                    merged_item['updated_at'] = datetime.utcnow().isoformat()
                    # Ensure status/stage reflect completion
                    merged_item['status'] = 'completed'
                    merged_item['stage'] = 'completed'
                    merged_item['progress'] = 100
                    item = merged_item
                    logger.info(f"Updating existing job record {analysis_id} with timestamp {timestamp}")
                else:
                    logger.warning(f"Job record {analysis_id} with timestamp {timestamp} not found, creating new record")
                    # Set default fields for new record
                    item['created_at'] = timestamp
                    item['updated_at'] = datetime.utcnow().isoformat()
                    item['status'] = 'completed'
                    item['stage'] = 'completed'
                    item['progress'] = 100
            except Exception as read_error:
                logger.warning(f"Could not read existing job record, proceeding with new record: {read_error}")
                # Set default fields for new record
                item['created_at'] = timestamp or datetime.utcnow().isoformat()
                item['updated_at'] = datetime.utcnow().isoformat()
                item['status'] = 'completed'
                item['stage'] = 'completed'
                item['progress'] = 100
        else:
            # New record (no timestamp provided) - set default fields
            item['created_at'] = datetime.utcnow().isoformat()
            item['updated_at'] = datetime.utcnow().isoformat()
        
        # Save to DynamoDB (put_item will create or replace entire item)
        table.put_item(Item=item)
        

        
        return {
            "success": True,
            "analysis_id": analysis_id,
            "conflicts_saved": len(conflicts)
        }
        
    except Exception as e:
        logger.error(f"Error saving analysis to DynamoDB: {str(e)}")
        return {
            "success": False,
            "error": str(e)
        }


def _save_and_upload_document(doc, bucket: str, s3_key: str, metadata: Dict[str, str]) -> bool:
    """
    Save python-docx document to memory and upload to S3.
    
    Args:
        doc: python-docx Document object
        bucket: S3 bucket name
        s3_key: S3 key for the uploaded document
        metadata: Dictionary of metadata to attach to S3 object
        
    Returns:
        Boolean indicating success
    """
    
    try:

        
        # Save document to memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Upload to S3 with metadata
        s3_client.put_object(
            Bucket=bucket,
            Key=s3_key,
            Body=buffer.getvalue(),
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            Metadata=metadata
        )
        

        return True
        
    except Exception as e:
        logger.error(f"Error saving and uploading document: {str(e)}")
        return False


def _cleanup_session_documents(session_id: str, user_id: str) -> Dict[str, Any]:
    """
    Clean up session reference documents using existing Lambda functions.
    
    Steps:
    1. List session reference documents
    2. Call delete_from_s3 Lambda function
    3. Call sync_knowledge_base Lambda function
    """
    try:

        
        # Step 1: List session reference documents
        session_s3_keys = _list_session_reference_documents(session_id, user_id)
        
        if not session_s3_keys:
            logger.info(f"No reference documents found for cleanup in session {session_id}")
            return {
                'success': True,
                'deleted_count': 0,
                'sync_triggered': False,
                'message': 'No documents to clean up'
            }
        
        logger.info(f"Found {len(session_s3_keys)} reference documents to delete")
        
        # Step 2: Delete documents using existing delete_from_s3 Lambda
        delete_result = _invoke_delete_lambda(session_s3_keys)
        
        if not delete_result.get('success'):
            return {
                'success': False,
                'error': f"Document deletion failed: {delete_result.get('error')}",
                'step_failed': 'delete',
                'found_documents': len(session_s3_keys)
            }
        
        deleted_count = delete_result.get('deleted_count', 0)
        logger.info(f"Successfully deleted {deleted_count} reference documents")
        
        # Step 3: Trigger knowledge base sync using existing sync Lambda
        sync_result = _invoke_sync_lambda()
        
        return {
            'success': True,
            'deleted_count': deleted_count,
            'sync_triggered': sync_result.get('success', False),
            'sync_job_id': sync_result.get('job_id'),
            'message': f"Cleanup completed: {deleted_count} documents deleted, sync triggered"
        }
        
    except Exception as e:
        logger.error(f"Session cleanup failed: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }

def _get_function_names() -> Dict[str, str]:
    """Get Lambda function names based on current function naming pattern."""
    current_function = os.environ.get('AWS_LAMBDA_FUNCTION_NAME', '')
    
    if current_function and 'stepfunctions-generateredline' in current_function:
        # Extract stack name dynamically from function name
        # Example: OneL-DV2-stepfunctions-generateredline -> OneL-DV2
        #          OneL-v2-stepfunctions-generateredline -> OneL-v2
        stack_name = current_function.replace('-stepfunctions-generateredline', '')
        
        return {
            'delete_function': f"{stack_name}-delete-from-s3",
            'sync_function': f"{stack_name}-sync-knowledge-base"
        }
    else:
        # Fallback: use stack name from constants if available, otherwise hardcoded default
        try:
            import sys
            import os as os_module
            _parent_dir = os_module.path.dirname(os_module.path.dirname(os_module.path.dirname(os_module.path.dirname(os_module.path.abspath(__file__)))))
            if _parent_dir not in sys.path:
                sys.path.insert(0, _parent_dir)
            import constants
            stack_name = constants.STACK_NAME
        except ImportError:
            # Final fallback: extract from current function name if available
            # This ensures we always use the correct stack name from the deployment
            current_function = os.environ.get('AWS_LAMBDA_FUNCTION_NAME', '')
            if current_function and '-' in current_function:
                # Extract stack name from function name pattern: <stack-name>-<function-type>-<function-name>
                # Example: OneL-DV2-stepfunctions-generateredline -> OneL-DV2
                parts = current_function.split('-')
                if len(parts) >= 2:
                    # Try to find where function type starts (stepfunctions, knowledge-management, etc.)
                    function_types = ['stepfunctions', 'knowledge', 'websocket', 'session', 'upload', 'retrieve', 'delete', 'sync', 'create', 'auth']
                    stack_parts = []
                    for part in parts:
                        if part in function_types:
                            break
                        stack_parts.append(part)
                    if stack_parts:
                        stack_name = '-'.join(stack_parts)
                    else:
                        # If we can't parse, use first two parts as fallback
                        stack_name = '-'.join(parts[:2])
                else:
                    raise ValueError(f"Cannot determine stack name from function name: {current_function}")
            else:
                raise ValueError("Cannot determine stack name: constants module not available and AWS_LAMBDA_FUNCTION_NAME not set")
        
        return {
            'delete_function': f'{stack_name}-delete-from-s3',
            'sync_function': f'{stack_name}-sync-knowledge-base'
        }


def _list_session_reference_documents(session_id: str, user_id: str) -> List[str]:
    """List all reference document S3 keys for a session."""
    try:
        bucket_name = os.environ.get('USER_DOCUMENTS_BUCKET')
        if not bucket_name:
            logger.error('USER_DOCUMENTS_BUCKET not configured')
            return []
        
        # Session reference docs are at: sessions/{user_id}/{session_id}/reference-docs/
        prefix = f"sessions/{user_id}/{session_id}/reference-docs/"
        
        logger.info(f"Listing documents with prefix: {prefix}")
        
        # List all objects with this prefix
        response = s3_client.list_objects_v2(Bucket=bucket_name, Prefix=prefix)
        
        if 'Contents' not in response:
            return []
        
        # Extract S3 keys
        s3_keys = [obj['Key'] for obj in response['Contents']]
        
        logger.info(f"Found {len(s3_keys)} reference documents: {s3_keys}")
        return s3_keys
        
    except Exception as e:
        logger.error(f"Error listing session documents: {str(e)}")
        return []


def _invoke_delete_lambda(s3_keys: List[str]) -> Dict[str, Any]:
    """Invoke the existing delete_from_s3 Lambda function."""
    try:
        lambda_client = boto3.client('lambda')
        
        function_names = _get_function_names()
        delete_function_name = function_names['delete_function']
        
        # Prepare payload for delete Lambda
        delete_payload = {
            'bucket_type': 'user_documents',
            's3_keys': s3_keys
        }
        
        logger.info(f"Invoking delete Lambda: {delete_function_name}")
        
        # Invoke delete Lambda synchronously to ensure completion
        response = lambda_client.invoke(
            FunctionName=delete_function_name,
            InvocationType='RequestResponse',  # Synchronous
            Payload=json.dumps(delete_payload)
        )
        
        # Parse response
        response_payload = json.loads(response['Payload'].read())
        
        if response_payload.get('statusCode') == 200:
            body = json.loads(response_payload.get('body', '{}'))
            logger.info(f"Delete Lambda succeeded: {body.get('message')}")
            return {
                'success': True,
                'deleted_count': body.get('deleted_count', 0),
                'details': body
            }
        else:
            error_msg = response_payload.get('body', 'Unknown error')
            logger.error(f"Delete Lambda failed: {error_msg}")
            return {
                'success': False,
                'error': error_msg
            }
            
    except Exception as e:
        logger.error(f"Error invoking delete Lambda: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }


def _invoke_sync_lambda() -> Dict[str, Any]:
    """Invoke the existing sync_knowledge_base Lambda function."""
    try:
        lambda_client = boto3.client('lambda')
        
        function_names = _get_function_names()
        sync_function_name = function_names['sync_function']
        
        # Prepare payload for sync Lambda
        sync_payload = {
            'action': 'start_sync',
            'data_source': 'user_documents',
            'triggered_by': 'session_cleanup'
        }
        
        logger.info(f"Invoking sync Lambda: {sync_function_name}")
        
        # Invoke sync Lambda asynchronously (sync can take time)
        response = lambda_client.invoke(
            FunctionName=sync_function_name,
            InvocationType='Event',  # Asynchronous
            Payload=json.dumps(sync_payload)
        )
        
        logger.info(f"Sync Lambda invoked successfully")
        
        return {
            'success': True,
            'function_invoked': sync_function_name,
            'message': 'Knowledge base sync triggered'
        }
        
    except Exception as e:
        logger.error(f"Error invoking sync Lambda: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }


 