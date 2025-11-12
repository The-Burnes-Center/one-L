"""
BACKUP: PyMuPDF-based PDF to DOCX conversion function
This is the original implementation using PyMuPDF + python-docx
Saved on: 2025-01-XX
To restore: Copy the _convert_pdf_to_docx_in_processing_bucket function from this file
"""

import logging
import io
import os
import re
from docx import Document
from docx.shared import RGBColor, Pt, Inches
import boto3

logger = logging.getLogger()
s3_client = boto3.client('s3')

def _convert_pdf_to_docx_in_processing_bucket_PYMUPDF_BACKUP(agent_bucket: str, pdf_s3_key: str) -> str:
    """
    [BACKUP] Convert a PDF stored in the processing bucket into a DOCX file with formatting preservation.
    Uses PyMuPDF + python-docx for conversion with enhanced formatting preservation.
    Returns the new DOCX S3 key on success.
    
    This is the ORIGINAL implementation - kept as backup for reverting if needed.
    
    Args:
        agent_bucket: S3 bucket name where PDF is stored
        pdf_s3_key: S3 key of the PDF file
        
    Returns:
        S3 key of the converted DOCX file
    """
    try:
        logger.info(f"PDF_TO_DOCX_START: Converting {pdf_s3_key} to DOCX with formatting preservation")
        
        # Download PDF from S3
        response = s3_client.get_object(Bucket=agent_bucket, Key=pdf_s3_key)
        pdf_bytes = response['Body'].read()
        logger.info(f"PDF_TO_DOCX: Downloaded PDF, size: {len(pdf_bytes)} bytes")
        
        # Use PyMuPDF + python-docx for conversion (enhanced with formatting preservation)
        logger.info("PDF_TO_DOCX: Using PyMuPDF + python-docx with enhanced formatting preservation")
        try:
            import fitz  # PyMuPDF
            
            docx_doc = Document()
            pdf_file = io.BytesIO(pdf_bytes)
            pdf = fitz.open(stream=pdf_file, filetype="pdf")
            
            for page_index in range(len(pdf)):
                page = pdf[page_index]
                try:
                    # Extract images from page first
                    try:
                        image_list = page.get_images()
                        if image_list:
                            logger.info(f"PDF_TO_DOCX: Found {len(image_list)} images on page {page_index + 1}")
                            for img_idx, img in enumerate(image_list):
                                try:
                                    xref = img[0]
                                    base_image = pdf.extract_image(xref)
                                    image_bytes = base_image["image"]
                                    
                                    if image_bytes:
                                        img_stream = io.BytesIO(image_bytes)
                                        para_img = docx_doc.add_paragraph()
                                        run_img = para_img.add_run()
                                        run_img.add_picture(img_stream, width=Inches(6))
                                        logger.info(f"PDF_TO_DOCX: Added image {img_idx + 1} to page {page_index + 1}")
                                except Exception as img_error:
                                    logger.warning(f"PDF_TO_DOCX: Error extracting image {img_idx + 1}: {img_error}")
                                    continue
                    except Exception as img_extract_error:
                        logger.debug(f"PDF_TO_DOCX: Image extraction error: {img_extract_error}")
                    
                    # Try to extract tables first (no page markers to preserve exact PDF structure)
                    try:
                        tables = page.find_tables()
                        if tables:
                            logger.info(f"PDF_TO_DOCX: Found {len(tables)} tables on page {page_index + 1}")
                            for table_idx, table in enumerate(tables):
                                try:
                                    table_data = table.extract()
                                    if table_data and len(table_data) > 0:
                                        docx_table = docx_doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
                                        docx_table.style = 'Light Grid Accent 1'
                                        
                                        for row_idx, row_data in enumerate(table_data):
                                            if row_idx < len(docx_table.rows):
                                                for col_idx, cell_data in enumerate(row_data):
                                                    if col_idx < len(docx_table.rows[row_idx].cells):
                                                        cell = docx_table.rows[row_idx].cells[col_idx]
                                                        cell.text = str(cell_data) if cell_data else ""
                                        logger.info(f"PDF_TO_DOCX: Added table {table_idx + 1} with {len(table_data)} rows")
                                except Exception as table_error:
                                    logger.warning(f"PDF_TO_DOCX: Error extracting table {table_idx + 1}: {table_error}")
                                    continue
                    except (AttributeError, Exception) as table_extract_error:
                        logger.debug(f"PDF_TO_DOCX: Table extraction not available: {table_extract_error}")
                    
                    # Extract text blocks with EXACT line-by-line preservation
                    # This ensures line breaks, alignment, and sentence boundaries match PDF exactly
                    text_dict = page.get_text("dict")
                    
                    for block in text_dict.get("blocks", []):
                        if "lines" in block:
                            # Process each line separately to preserve exact line breaks
                            for line_idx, line in enumerate(block["lines"]):
                                # Collect text from this line only (preserve line boundaries)
                                line_text = ""
                                spans_data = []
                                
                                for span in line.get("spans", []):
                                    # Preserve exact text including spaces (don't strip)
                                    text = span.get("text", "")
                                    if text:
                                        flags = span.get("flags", 0)
                                        font_size = span.get("size", 11)
                                        font_color = span.get("color", 0)
                                        font_name = span.get("font", "")
                                        
                                        line_text += text  # Preserve exact text with spaces
                                        spans_data.append({
                                            'text': text,
                                            'bold': bool(flags & 16),
                                            'italic': bool(flags & 2),
                                            'size': font_size,
                                            'color': font_color,
                                            'font': font_name
                                        })
                                
                                # Skip empty lines
                                if not line_text.strip():
                                    continue
                                
                                # Detect numbered list patterns on this line
                                line_text_stripped = line_text.strip()
                                is_numbered_list = False
                                list_style = None
                                
                                if re.match(r'^\d+[a-z]\b', line_text_stripped, re.IGNORECASE):
                                    is_numbered_list = True
                                    list_style = 'List Number 2'
                                elif re.match(r'^[\d]+[\.\)]', line_text_stripped):
                                    is_numbered_list = True
                                    list_style = 'List Number'
                                elif re.match(r'^[a-z][\.\)]', line_text_stripped, re.IGNORECASE):
                                    is_numbered_list = True
                                    list_style = 'List Bullet 2'
                                elif re.match(r'^\([a-z0-9]+\)', line_text_stripped, re.IGNORECASE):
                                    is_numbered_list = True
                                    list_style = 'List Bullet 2'
                                elif re.search(r'\b\d+[a-z]\b', line_text_stripped, re.IGNORECASE) and len(line_text_stripped) < 100:
                                    is_numbered_list = True
                                    list_style = 'List Number 2'
                                
                                # Create ONE paragraph per line to preserve exact line breaks
                                if is_numbered_list:
                                    para = docx_doc.add_paragraph(style=list_style)
                                else:
                                    para = docx_doc.add_paragraph()
                                
                                # Add each span with preserved formatting and exact spacing
                                for span_idx, span_data in enumerate(spans_data):
                                    run = para.add_run(span_data['text'])  # Preserve exact text including spaces
                                    
                                    if span_data['bold']:
                                        run.font.bold = True
                                    if span_data['italic']:
                                        run.font.italic = True
                                    
                                    try:
                                        if span_data['size'] > 0:
                                            run.font.size = Pt(span_data['size'])
                                    except:
                                        pass
                                    
                                    try:
                                        color_int = span_data['color']
                                        r = (color_int >> 16) & 0xFF
                                        g = (color_int >> 8) & 0xFF
                                        b = color_int & 0xFF
                                        if not (r == 0 and g == 0 and b == 0):
                                            run.font.color.rgb = RGBColor(r, g, b)
                                    except:
                                        pass
                                    
                                    try:
                                        if span_data['font']:
                                            font_map = {
                                                'Arial': 'Arial',
                                                'ArialMT': 'Arial',
                                                'Arial-BoldMT': 'Arial',
                                                'Arial-ItalicMT': 'Arial',
                                                'Helvetica': 'Arial',
                                                'Helvetica-Bold': 'Arial',
                                                'Helvetica-Oblique': 'Arial',
                                                'Times-Roman': 'Times New Roman',
                                                'TimesNewRomanPSMT': 'Times New Roman',
                                                'TimesNewRomanPS-BoldMT': 'Times New Roman',
                                                'TimesNewRomanPS-ItalicMT': 'Times New Roman',
                                                'Courier': 'Courier New',
                                                'CourierNew': 'Courier New',
                                                'CourierNewPSMT': 'Courier New',
                                                'Calibri': 'Calibri',
                                                'Calibri-Bold': 'Calibri',
                                                'Calibri-Italic': 'Calibri',
                                            }
                                            if span_data['font'] not in ['Times-Roman']:
                                                font_name = font_map.get(span_data['font'], span_data['font'])
                                                try:
                                                    run.font.name = font_name
                                                except Exception as font_err:
                                                    logger.debug(f"PDF_TO_DOCX: Could not set font {font_name}: {font_err}")
                                                    pass
                                    except Exception as font_error:
                                        logger.debug(f"PDF_TO_DOCX: Font processing error: {font_error}")
                                        pass
                                
                                # Each line becomes a separate paragraph - preserves exact line breaks
                                # No need to add spaces between spans since we preserve exact text
                except Exception as page_error:
                    logger.warning(f"PDF_TO_DOCX: Error processing page {page_index + 1}: {page_error}")
                    continue
            
            pdf.close()
            
            # Save DOCX to memory
            docx_bytes_io = io.BytesIO()
            docx_doc.save(docx_bytes_io)
            docx_bytes_io.seek(0)
            docx_bytes = docx_bytes_io.read()
            
            # Upload DOCX to S3
            new_key = pdf_s3_key.rsplit('.', 1)[0] + '.docx'
            s3_client.put_object(
                Bucket=agent_bucket,
                Key=new_key,
                Body=docx_bytes,
                ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            logger.info(f"PDF_TO_DOCX_SUCCESS: Converted to {new_key} (enhanced formatting: fonts, sizes, colors, styles, images, lists preserved)")
            return new_key
                
        except Exception as fallback_error:
            logger.error(f"PDF_TO_DOCX_FALLBACK_FAILED: {str(fallback_error)}")
            raise Exception(f"PDF to DOCX conversion failed: {str(fallback_error)}")
            
    except Exception as e:
        logger.error(f"PDF_TO_DOCX_ERROR: Failed to convert PDF to DOCX: {str(e)}")
        raise Exception(f"Failed to convert PDF to DOCX: {str(e)}")

